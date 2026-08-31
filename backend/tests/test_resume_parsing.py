"""Text-extraction tests for the resume parser (Track B).

Covers DOCX (python-docx, including tables), PDF (pypdf), format sniffing when
the extension is missing, and — most importantly — that empty / corrupt /
wrong-format bytes degrade to "" instead of raising into the Celery task.

The DOCX fixture is generated in-process with python-docx so there is no binary
blob checked into the repo. The PDF fixture is a tiny hand-built one-page PDF.
"""
import io
import uuid
from types import SimpleNamespace

import pytest

from app.services import resume_parsing
from app.services.resume_parsing import extract_text


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ── DOCX ─────────────────────────────────────────────────────────────────────

def test_extract_docx_paragraphs():
    data = _make_docx(["Jane Doe", "Senior Engineer", "Python, FastAPI, Postgres"])
    text = extract_text(data, "resume.docx")
    assert "Jane Doe" in text
    assert "FastAPI" in text


def test_extract_docx_includes_table_cells():
    data = _make_docx(
        ["Experience"],
        table_rows=[["Company", "Role"], ["Acme Corp", "Lead Developer"]],
    )
    text = extract_text(data, "resume.docx")
    assert "Acme Corp" in text
    assert "Lead Developer" in text


def test_extract_docx_empty_document_returns_empty_string():
    data = _make_docx([])
    assert extract_text(data, "empty.docx") == ""


def test_extract_docx_sniffed_without_extension():
    # Cloudinary raw URLs can drop the extension — bytes start with "PK" (zip).
    data = _make_docx(["No extension here"])
    assert "No extension here" in extract_text(data, "download")


# ── PDF ──────────────────────────────────────────────────────────────────────

def _make_pdf(body_text: str) -> bytes:
    """Build a valid single-page PDF with a correct cross-reference table so
    strict pypdf parsers accept it (no PDF-writer library is installed)."""
    stream = f"BT /F1 24 Tf 72 700 Td ({body_text}) Tj ET".encode("latin-1")
    objs = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>",
        b"<</Length %d>>\nstream\n%s\nendstream" % (len(stream), stream),
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + obj + b"\nendobj\n"
    xref_pos = len(out)
    out += b"xref\n0 %d\n" % (len(objs) + 1)
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<</Size %d/Root 1 0 R>>\n" % (len(objs) + 1)
    out += b"startxref\n%d\n%%%%EOF" % xref_pos
    return bytes(out)


def test_extract_pdf_text():
    text = extract_text(_make_pdf("Hello Resume"), "resume.pdf")
    assert "Hello Resume" in text


def test_extract_pdf_sniffed_without_extension():
    text = extract_text(_make_pdf("Hello Resume"), "download")
    assert "Hello Resume" in text


# ── Garbage / empty / never-raises ───────────────────────────────────────────

@pytest.mark.parametrize(
    "data,filename",
    [
        (b"", "resume.pdf"),
        (b"", "resume.docx"),
        (b"", "unknown"),
        (b"not a real document at all", "resume.pdf"),
        (b"not a real document at all", "resume.docx"),
        (b"\x00\x01\x02\x03garbage", "resume.docx"),
        (b"random bytes with no magic", "unknown"),
    ],
)
def test_extract_text_never_raises_on_garbage(data, filename):
    # The whole point: a corrupt/empty/unknown file yields "" and never raises
    # into the Celery task.
    assert extract_text(data, filename) == ""


def test_extract_text_none_filename():
    assert extract_text(b"", None) == ""


# ── Structured extraction short-circuits on empty text (no LLM call) ─────────

@pytest.mark.asyncio
async def test_extract_structured_fields_empty_text_no_llm(monkeypatch):
    called = False

    async def _boom(*a, **k):
        nonlocal called
        called = True
        raise AssertionError("LLM must not be called for empty text")

    monkeypatch.setattr(resume_parsing.llm_router, "chat_completion", _boom)
    result = await resume_parsing.extract_structured_fields("   ")
    assert called is False
    assert result == {
        "skills": [],
        "total_experience_years": None,
        "education": [],
        "employment_history": [],
    }


@pytest.mark.asyncio
async def test_extract_structured_fields_non_json_degrades_to_empty(monkeypatch):
    async def _prose(*a, **k):
        return "Sorry, I can't help with that."

    monkeypatch.setattr(resume_parsing.llm_router, "chat_completion", _prose)
    result = await resume_parsing.extract_structured_fields("some real resume text")
    assert result == {
        "skills": [],
        "total_experience_years": None,
        "education": [],
        "employment_history": [],
    }


@pytest.mark.asyncio
async def test_extract_structured_fields_valid_json(monkeypatch):
    async def _ok(*a, **k):
        return (
            '{"skills": ["Python"], "total_experience_years": 5, '
            '"education": [{"degree": "BE", "institution": "X", "year": 2015}], '
            '"employment_history": [{"company": "Acme", "title": "Dev", '
            '"start": "2018-01", "end": null}]}'
        )

    monkeypatch.setattr(resume_parsing.llm_router, "chat_completion", _ok)
    result = await resume_parsing.extract_structured_fields("real text")
    assert result["skills"] == ["Python"]
    assert result["total_experience_years"] == 5
    assert result["employment_history"][0]["company"] == "Acme"


@pytest.mark.asyncio
async def test_parse_resume_indexes_text_when_llm_is_unavailable(monkeypatch):
    profile = SimpleNamespace(
        id=uuid.uuid4(),
        candidate_id=uuid.uuid4(),
        resume_text="Python FastAPI PostgreSQL",
        parsed_fields_json=None,
        embedding=None,
    )

    class _Session:
        committed = False

        async def get(self, _model, _profile_id):
            return profile

        async def commit(self):
            self.committed = True

    async def _llm_down(*_args, **_kwargs):
        raise resume_parsing.llm_router.LLMUnavailableError("quota exhausted")

    async def _embed(texts):
        assert texts == ["Python FastAPI PostgreSQL"]
        return [[0.25] * 1024]

    session = _Session()
    monkeypatch.setattr(resume_parsing, "extract_structured_fields", _llm_down)
    monkeypatch.setattr(resume_parsing, "embed", _embed)

    await resume_parsing.parse_resume(session, uuid.uuid4())

    assert session.committed is True
    assert profile.parsed_fields_json == resume_parsing._EMPTY_PARSED_FIELDS
    assert profile.embedding == [0.25] * 1024
