"""Candidate sourcing, review-screen, decisions, pipeline status and
interview scheduling (FR-4.3/4.4, FR-7.x, FR-8.x)."""
import html
import io
import uuid
from urllib.parse import urlencode
from datetime import datetime, timezone

from docx import Document
from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    Request,
    UploadFile,
    status,
)
from fastapi.responses import HTMLResponse, RedirectResponse, Response
import jwt
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import CurrentUser, get_tenant_db, require_capability
from app.models.candidate import (
    Candidate,
    CandidateTeamReview,
    Interview,
    JobCandidateLink,
    PipelineStatusEntry,
    Profile,
    SOURCE_TYPE_SOURCED,
    VerificationRequest,
    source_type_label,
)
from app.models.enums import LinkSource, PipelineStatus
from app.models.job import Job
from app.models.tenant import Tenant
from app.models.user import User
from app.schemas.candidates import (
    CandidateOut,
    DecisionIn,
    GrantAccessOut,
    InterviewIn,
    InterviewOut,
    JobLinksOut,
    LinkArchiveOut,
    LinkOut,
    ProfileOut,
    RankingCommentsOut,
    StatusIn,
    StatusOut,
    TeamReviewIn,
    TeamReviewOut,
    TeamReviewRewriteIn,
    TeamReviewRewriteOut,
    TeamReviewsOut,
    UploadResumeOut,
    VerificationRequestSummary,
)
from app.services import capabilities as caps
from app.services import rbac
from app.services import team_review
from app.services.audit import audit
from app.services.matching import client_breakdown, ranking_payload
from app.services import llm_router
from app.services.resume_storage import (
    ALLOWED_RESUME_CONTENT_TYPES,
    ALLOWED_RESUME_EXTENSIONS as ALLOWED_RESUME_EXTS,
    MAX_RESUME_BYTES,
    apply_resume_asset,
    fetch_resume_bytes,
    read_validated_resume,
    ResumeStorageError,
    store_resume,
)
from app.services.resume_access import issue_resume_token, verify_resume_token
from app.workers.celery_app import celery_app

router = APIRouter()

#: Statuses that move an application FORWARD. `offer_extended` is the current
#: name for `offered` (migration 0018 kept both valid), and omitting it meant a
#: fresh-sourced candidate advanced under the new name was not treated as
#: progressed at all.
FORWARD_STATUSES = {
    PipelineStatus.shortlisted,
    PipelineStatus.interview_scheduled,
    PipelineStatus.interview_completed,
    PipelineStatus.offered,
    PipelineStatus.offer_extended,
    PipelineStatus.joined,
}


async def _get_link(
    session: AsyncSession, user: CurrentUser, link_id: uuid.UUID
) -> JobCandidateLink:
    link = await session.get(JobCandidateLink, link_id)
    if link is None or link.tenant_id != user.tenant_id:  # defense in depth
        raise HTTPException(status_code=404, detail="Link not found")
    return link


async def _latest_status(
    session: AsyncSession, link_ids: list[uuid.UUID]
) -> dict[uuid.UUID, PipelineStatusEntry]:
    if not link_ids:
        return {}
    rows = (
        await session.execute(
            select(PipelineStatusEntry)
            .where(PipelineStatusEntry.job_candidate_link_id.in_(link_ids))
            .order_by(PipelineStatusEntry.at)
        )
    ).scalars().all()
    latest: dict[uuid.UUID, PipelineStatusEntry] = {}
    for row in rows:  # ordered ascending, so the last write wins
        latest[row.job_candidate_link_id] = row
    return latest


@router.post(
    "/jobs/{job_id}/upload-resume",
    response_model=UploadResumeOut,
    status_code=status.HTTP_201_CREATED,
)
async def upload_resume(
    job_id: uuid.UUID,
    file: UploadFile = File(...),
    email: str = Form(...),
    full_name: str | None = Form(default=None),
    phone: str | None = Form(default=None),
    user: CurrentUser = Depends(require_capability(caps.UPLOAD_RESUMES)),
    session: AsyncSession = Depends(get_tenant_db),
) -> UploadResumeOut:
    """Recruiter uploads a freshly sourced resume (FR-4.3): creates
    candidate + profile + link (source=fresh) and enqueues parsing."""
    job = await session.get(Job, job_id)
    if job is None or job.tenant_id != user.tenant_id:
        raise HTTPException(status_code=404, detail="Job not found")
    if job.ratified_at is None:
        # ASSUMPTION: sourcing starts once the job has reached HR (FR-3.4).
        raise HTTPException(status_code=409, detail="Job is not ratified yet")

    candidate = (
        await session.execute(select(Candidate).where(Candidate.email == email))
    ).scalars().first()
    if candidate is None:
        candidate = Candidate(
            tenant_id=user.tenant_id, email=email, full_name=full_name, phone=phone
        )
        session.add(candidate)
        await session.flush()

    dup = (
        await session.execute(
            select(JobCandidateLink).where(
                JobCandidateLink.job_id == job.id,
                JobCandidateLink.candidate_id == candidate.id,
            )
        )
    ).scalars().first()
    if dup is not None:
        raise HTTPException(status_code=409, detail="Candidate is already linked to this job")

    asset = await store_resume(file)
    profile = Profile(candidate_id=candidate.id, source_tenant_id=user.tenant_id)
    apply_resume_asset(profile, asset)
    session.add(profile)
    await session.flush()

    link = JobCandidateLink(
        tenant_id=user.tenant_id, job_id=job.id, candidate_id=candidate.id,
        profile_id=profile.id, source=LinkSource.fresh,
        # ASSUMPTION (2026-07-28): a recruiter uploading ONE resume they found
        # elsewhere is `sourced`, not `databank`. Databank is specifically the
        # bulk upload (POST /jobs/{id}/candidates/databank, up to 25 files);
        # this single-file route predates it and describes a candidate the
        # recruiter procured from outside ReadyPick.
        source_type=SOURCE_TYPE_SOURCED,
    )
    session.add(link)
    await session.flush()

    celery_app.send_task("pickready.parse_resume", args=[str(profile.id)])
    await audit(session, tenant_id=user.tenant_id, actor_user_id=user.user_id,
                action="resume_uploaded", target_type="profile", target_id=profile.id,
                metadata={"job_id": str(job.id), "candidate_id": str(candidate.id),
                          "resume_public_id": asset.public_id, "resume_sha256": asset.sha256})
    return UploadResumeOut(
        candidate_id=candidate.id, profile_id=profile.id, link_id=link.id,
        source=LinkSource.fresh, resume_public_id=asset.public_id, resume_url=asset.secure_url,
    )


@router.get("/{candidate_id}/profile", response_model=ProfileOut)
async def get_profile(
    candidate_id: uuid.UUID,
    profile_id: uuid.UUID | None = Query(default=None),
    user: CurrentUser = Depends(require_capability(caps.VIEW_REVIEW_SCREEN)),
    session: AsyncSession = Depends(get_tenant_db),
) -> ProfileOut:
    """Full Profile for the HR Review Screen (FR-7.1/7.2).

    # ASSUMPTION: full access to any profile is derived from the HR-exclusive
    # SEND_OUTREACH capability; holders of only VIEW_REVIEW_SCREEN (Hiring
    # Managers) can read a candidate solely when HR has granted access on at
    # least one of that candidate's links (FR-8.1) — capability + data, no
    # role branch.
    """
    candidate = await session.get(Candidate, candidate_id)
    if candidate is None:
        raise HTTPException(status_code=404, detail="Candidate not found")

    full_access = await rbac.has_capability(
        session, user.tenant_id, user.role, caps.SEND_OUTREACH
    )
    if not full_access:
        granted = (
            await session.execute(
                select(JobCandidateLink).where(
                    JobCandidateLink.candidate_id == candidate_id,
                    JobCandidateLink.tenant_id == user.tenant_id,
                    JobCandidateLink.hm_access_granted.is_(True),
                )
            )
        ).scalars().first()
        if granted is None:
            raise HTTPException(status_code=403, detail="Profile access not granted")

    profile_query = select(Profile).where(Profile.candidate_id == candidate_id)
    if profile_id is not None:
        profile_query = profile_query.where(Profile.id == profile_id)
    else:
        profile_query = profile_query.order_by(Profile.created_at.desc())
    profile = (await session.execute(profile_query)).scalars().first()
    if profile is None:
        raise HTTPException(status_code=404, detail="No profile for this candidate")
    if profile_id is not None:
        # EXISTENCE check only. The same (candidate, tenant, profile) triple
        # legitimately appears on many rows — one per job the candidate is
        # linked to in this tenant, and a reused resume keeps the same profile
        # — so `scalar_one_or_none()` here raised MultipleResultsFound (500)
        # for any candidate linked to more than one job, which the review
        # screen surfaced as "Could not load this candidate's profile".
        scoped_profile_link = (
            await session.execute(
                select(JobCandidateLink.id)
                .where(
                    JobCandidateLink.candidate_id == candidate_id,
                    JobCandidateLink.tenant_id == user.tenant_id,
                    JobCandidateLink.profile_id == profile_id,
                )
                .limit(1)
            )
        ).scalars().first()
        if scoped_profile_link is None:
            raise HTTPException(status_code=404, detail="No profile for this candidate")

    vrs = (
        await session.execute(
            select(VerificationRequest)
            .where(VerificationRequest.profile_id == profile.id)
            .order_by(VerificationRequest.employer_seq)
        )
    ).scalars().all()
    return ProfileOut(
        id=profile.id,
        candidate=CandidateOut.model_validate(candidate),
        resume_url=profile.resume_url,
        resume_public_id=profile.resume_public_id,
        resume_original_filename=profile.resume_original_filename,
        resume_mime_type=profile.resume_mime_type,
        resume_size_bytes=profile.resume_size_bytes,
        resume_uploaded_at=profile.resume_uploaded_at,
        aspects_json=profile.aspects_json,
        parsed_fields_json=profile.parsed_fields_json,
        aspects_completed_at=profile.aspects_completed_at,
        verification_requests=[VerificationRequestSummary.model_validate(v) for v in vrs],
    )


@router.get("/{candidate_id}/project-evidence")
async def get_project_evidence(
    candidate_id: uuid.UUID,
    user: CurrentUser = Depends(require_capability(caps.VIEW_REVIEW_SCREEN)),
    session: AsyncSession = Depends(get_tenant_db),
) -> dict:
    """Derived Project Evidence for the review screen.

    Same access rules as the full profile (VIEW_REVIEW_SCREEN, with the
    HM-grant gate for non-HR holders), plus one stricter check: the candidate
    must be linked to a job in THIS tenant, and a candidate who is not answers
    404 rather than 403 (RBAC: a cross-tenant read must not confirm
    existence).

    Everything returned is DERIVED intelligence: candidate claims labelled as
    claims, observed evidence labelled as observed, strength as a WORD. There
    is no original file to offer, because none is retained.
    """
    from app.services.projects import context as project_context

    candidate = await session.get(Candidate, candidate_id)
    if candidate is None:
        raise HTTPException(status_code=404, detail="Candidate not found")
    linked = (
        await session.execute(
            select(JobCandidateLink.id)
            .where(
                JobCandidateLink.candidate_id == candidate_id,
                JobCandidateLink.tenant_id == user.tenant_id,
            )
            .limit(1)
        )
    ).scalars().first()
    if linked is None:
        raise HTTPException(status_code=404, detail="Candidate not found")
    full_access = await rbac.has_capability(
        session, user.tenant_id, user.role, caps.SEND_OUTREACH
    )
    if not full_access:
        granted = (
            await session.execute(
                select(JobCandidateLink).where(
                    JobCandidateLink.candidate_id == candidate_id,
                    JobCandidateLink.tenant_id == user.tenant_id,
                    JobCandidateLink.hm_access_granted.is_(True),
                )
            )
        ).scalars().first()
        if granted is None:
            raise HTTPException(status_code=403, detail="Profile access not granted")
    return {"projects": await project_context.recruiter_views(session, candidate_id)}


@router.get("/profiles/{profile_id}/resume-preview", response_class=HTMLResponse)
async def preview_resume(
    profile_id: uuid.UUID,
    user: CurrentUser = Depends(require_capability(caps.VIEW_REVIEW_SCREEN)),
    session: AsyncSession = Depends(get_tenant_db),
) -> HTMLResponse:
    """Render a DOCX resume inside ReadyPick as safe, monochrome HTML.

    Browsers cannot natively display Word documents. The server downloads only
    the trusted private asset already stored on the profile, extracts
    paragraphs and tables with python-docx, HTML-escapes every value, and
    returns a same-app preview. PDF/image previews continue to use the browser's
    native viewer.
    """
    profile = await session.get(Profile, profile_id)
    if profile is None or not profile.resume_url:
        raise HTTPException(status_code=404, detail="Resume not found")
    link = (
        await session.execute(
            select(JobCandidateLink).where(
                JobCandidateLink.profile_id == profile.id,
                JobCandidateLink.tenant_id == user.tenant_id,
            )
        )
    ).scalars().first()
    if link is None:
        raise HTTPException(status_code=404, detail="Resume not found")
    full_access = await rbac.has_capability(
        session, user.tenant_id, user.role, caps.SEND_OUTREACH
    )
    if not full_access and not link.hm_access_granted:
        raise HTTPException(status_code=403, detail="Profile access not granted")

    try:
        resume_bytes = await fetch_resume_bytes(profile)
    except ResumeStorageError as exc:
        raise HTTPException(status_code=502, detail="Resume could not be loaded") from exc
    try:
        document = Document(io.BytesIO(resume_bytes))
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Resume is not a readable DOCX file") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text_value = paragraph.text.strip()
        if not text_value:
            continue
        safe = html.escape(text_value)
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            blocks.append(f"<h2>{safe}</h2>")
        else:
            blocks.append(f"<p>{safe}</p>")
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = "".join(
                f"<td>{html.escape(cell.text.strip())}</td>" for cell in row.cells
            )
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            blocks.append(f"<table>{''.join(rows)}</table>")

    filename = html.escape(profile.resume_original_filename or "Resume")
    body = "\n".join(blocks) or "<p>No readable text was found in this document.</p>"
    return HTMLResponse(
        content=f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width">
<title>{filename}</title><style>
html{{background:#f5f5f5;color:#111;font-family:Arial,sans-serif}}
body{{box-sizing:border-box;max-width:850px;margin:24px auto;padding:56px;background:#fff;
box-shadow:0 1px 8px rgba(0,0,0,.14);line-height:1.55}}
h2{{font-size:1.1rem;margin:1.5rem 0 .5rem;border-bottom:1px solid #bbb;padding-bottom:.25rem}}
p{{margin:.55rem 0;white-space:pre-wrap}}table{{width:100%;border-collapse:collapse;margin:1rem 0}}
td{{border:1px solid #bbb;padding:.45rem;vertical-align:top}}
@media(max-width:700px){{body{{margin:0;padding:24px;box-shadow:none}}}}
</style></head><body>{body}</body></html>"""
    )


@router.get("/profiles/{profile_id}/resume-file")
async def resume_file(
    request: Request,
    profile_id: uuid.UUID,
    download: bool = Query(default=False),
    access_token: str | None = Query(default=None),
    user: CurrentUser = Depends(require_capability(caps.VIEW_REVIEW_SCREEN)),
    session: AsyncSession = Depends(get_tenant_db),
) -> Response:
    """Proxy an authorized resume without exposing a raw storage URL.

    The token round trip below redirects to THIS route. The target is taken
    from `request.url.path`, never written out by hand: this router is mounted
    at `/api/v1` only, and a hardcoded `/api/v2/candidates/...` here meant every
    resume view and every download 307ed to a path that does not exist and
    404ed. It read as a broken file (or, since only PDFs take this path while
    DOCX goes to `resume-preview`, as a broken FORMAT) rather than as a broken
    URL. Deriving the path means the mount and the redirect cannot drift again.
    """
    profile = await session.get(Profile, profile_id)
    if profile is None or not profile.resume_url:
        raise HTTPException(status_code=404, detail="Resume not found")
    link = (
        await session.execute(
            select(JobCandidateLink).where(
                JobCandidateLink.profile_id == profile.id,
                JobCandidateLink.tenant_id == user.tenant_id,
            )
        )
    ).scalars().first()
    if link is None:
        raise HTTPException(status_code=404, detail="Resume not found")
    full_access = await rbac.has_capability(
        session, user.tenant_id, user.role, caps.SEND_OUTREACH
    )
    if not full_access and not link.hm_access_granted:
        raise HTTPException(status_code=403, detail="Profile access not granted")
    if access_token is None:
        query = urlencode(
            {
                "download": str(download).lower(),
                "access_token": issue_resume_token(profile_id, user.tenant_id),
            }
        )
        return RedirectResponse(
            url=f"{request.url.path}?{query}",
            status_code=status.HTTP_307_TEMPORARY_REDIRECT,
            headers={"Cache-Control": "private, no-store"},
        )
    try:
        verify_resume_token(access_token, profile_id, user.tenant_id)
    except jwt.PyJWTError as exc:
        raise HTTPException(status_code=403, detail="Resume link is invalid or expired") from exc
    try:
        content = await fetch_resume_bytes(profile)
    except ResumeStorageError as exc:
        raise HTTPException(status_code=502, detail="Resume could not be loaded") from exc
    filename = (profile.resume_original_filename or "resume").replace('"', "")
    disposition = "attachment" if download else "inline"
    return Response(
        content=content,
        media_type=profile.resume_mime_type or "application/octet-stream",
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Cache-Control": "private, no-store",
            "X-Content-Type-Options": "nosniff",
        },
    )


@router.get("/{candidate_id}/ranking", response_model=RankingCommentsOut)
async def get_candidate_ranking(
    candidate_id: uuid.UUID,
    job_id: uuid.UUID | None = Query(default=None),
    user: CurrentUser = Depends(require_capability(caps.VIEW_REVIEW_SCREEN)),
    session: AsyncSession = Depends(get_tenant_db),
) -> RankingCommentsOut:
    """Return only the five human-readable ranking comments for a candidate.

    A candidate may be ranked for several jobs. Callers should pass `job_id`;
    when omitted, the newest ranking in the current tenant is returned for
    backwards compatibility with the path-only contract.
    """
    filters = [
        JobCandidateLink.candidate_id == candidate_id,
        JobCandidateLink.tenant_id == user.tenant_id,
    ]
    if job_id is not None:
        filters.append(JobCandidateLink.job_id == job_id)
    link = (
        await session.execute(
            select(JobCandidateLink)
            .where(*filters)
            .order_by(JobCandidateLink.created_at.desc())
        )
    ).scalars().first()
    if link is None:
        raise HTTPException(status_code=404, detail="Ranking not found")
    payload = ranking_payload(link.match_breakdown_json)
    return RankingCommentsOut(
        skills_match_comment=payload["skills_match_comment"],
        experience_comment=payload["experience_comment"],
        role_alignment_comment=payload["role_alignment_comment"],
        education_comment=payload["education_comment"],
        overall_comment=payload["overall_comment"],
    )


@router.post("/links/{link_id}/grant-access", response_model=GrantAccessOut)
async def grant_access(
    link_id: uuid.UUID,
    user: CurrentUser = Depends(require_capability(caps.SEND_OUTREACH)),
    session: AsyncSession = Depends(get_tenant_db),
) -> GrantAccessOut:
    """HR grants Hiring Manager access to a reviewed profile (FR-8.1).
    # ASSUMPTION: gated by SEND_OUTREACH — the HR-exclusive candidate-management
    # capability — since the PRD matrix has no dedicated grant capability."""
    link = await _get_link(session, user, link_id)
    link.hm_access_granted = True
    await session.flush()
    await audit(session, tenant_id=user.tenant_id, actor_user_id=user.user_id,
                action="hm_access_granted", target_type="job_candidate_link",
                target_id=link.id)
    return GrantAccessOut(link_id=link.id)


@router.delete("/links/{link_id}", response_model=LinkArchiveOut)
async def archive_candidate_application(
    link_id: uuid.UUID,
    user: CurrentUser = Depends(require_capability(caps.UPDATE_PIPELINE_STATUS)),
    session: AsyncSession = Depends(get_tenant_db),
) -> LinkArchiveOut:
    """Archive this job application without deleting the shared candidate."""
    link = await _get_link(session, user, link_id)
    if link.archived_at is None:
        link.archived_at = datetime.now(timezone.utc)
        await session.flush()
        await audit(
            session,
            tenant_id=user.tenant_id,
            actor_user_id=user.user_id,
            action="candidate_application_archived",
            target_type="job_candidate_link",
            target_id=link.id,
        )
    return LinkArchiveOut(link_id=link.id, archived=True)


@router.post("/links/{link_id}/restore", response_model=LinkArchiveOut)
async def restore_candidate_application(
    link_id: uuid.UUID,
    user: CurrentUser = Depends(require_capability(caps.UPDATE_PIPELINE_STATUS)),
    session: AsyncSession = Depends(get_tenant_db),
) -> LinkArchiveOut:
    link = await _get_link(session, user, link_id)
    if link.archived_at is not None:
        link.archived_at = None
        await session.flush()
        await audit(
            session,
            tenant_id=user.tenant_id,
            actor_user_id=user.user_id,
            action="candidate_application_restored",
            target_type="job_candidate_link",
            target_id=link.id,
        )
    return LinkArchiveOut(link_id=link.id, archived=False)


@router.post("/links/{link_id}/decision", response_model=StatusOut)
async def decide_profile(
    link_id: uuid.UUID,
    body: DecisionIn,  # hold without remarks -> 422 (schema validator)
    user: CurrentUser = Depends(require_capability(caps.DECIDE_PROFILE)),
    session: AsyncSession = Depends(get_tenant_db),
) -> StatusOut:
    """Hiring Manager decision: Rejected / Shortlisted / Hold (FR-8.2)."""
    link = await _get_link(session, user, link_id)
    if not link.hm_access_granted:
        raise HTTPException(status_code=403, detail="Profile access not granted (FR-8.1)")

    entry = PipelineStatusEntry(
        tenant_id=user.tenant_id, job_candidate_link_id=link.id,
        status=PipelineStatus(body.status), remarks=body.remarks, set_by=user.user_id,
    )
    session.add(entry)
    await session.flush()
    await audit(session, tenant_id=user.tenant_id, actor_user_id=user.user_id,
                action="profile_decision", target_type="job_candidate_link",
                target_id=link.id, metadata={"status": body.status, "remarks": body.remarks})
    return StatusOut(link_id=link.id, status=entry.status, remarks=entry.remarks,
                     at=entry.at or datetime.now(timezone.utc))


#: Display and tie-break order only; nothing scores these. Read from the one
#: place the vocabulary is defined so a fourth verdict cannot be added here and
#: nowhere else.
_TEAM_RATING_ORDER = list(team_review.VERDICTS)


def _clean_review_text(value: str) -> str:
    text_value = " ".join(value.strip().split())
    if not text_value:
        return text_value
    text_value = text_value[0].upper() + text_value[1:]
    if text_value[-1] not in ".!?":
        text_value += "."
    return text_value.replace(chr(8212), " - ")


async def _team_reviews_out(
    session: AsyncSession,
    link_id: uuid.UUID,
    current_user_id: uuid.UUID,
) -> TeamReviewsOut:
    rows = (
        await session.execute(
            select(CandidateTeamReview, User.full_name, User.email)
            .join(User, User.id == CandidateTeamReview.reviewer_user_id)
            .where(CandidateTeamReview.job_candidate_link_id == link_id)
            .order_by(CandidateTeamReview.updated_at.desc(), CandidateTeamReview.id)
        )
    ).all()
    reviews = [
        TeamReviewOut(
            id=review.id,
            reviewer_user_id=review.reviewer_user_id,
            reviewer_name=full_name or email or "Team member",
            rating=review.rating,
            remarks=review.remarks,
            ai_rewritten_remarks=review.ai_rewritten_remarks,
            is_current_user=review.reviewer_user_id == current_user_id,
            created_at=review.created_at,
            updated_at=review.updated_at,
        )
        for review, full_name, email in rows
    ]
    if not reviews:
        return TeamReviewsOut(
            reviews=[], overall_rating=None, overall_remarks=None, review_count=0
        )

    ordered = sorted(_TEAM_RATING_ORDER.index(review.rating) for review in reviews)
    overall_rating = _TEAM_RATING_ORDER[ordered[(len(ordered) - 1) // 2]]
    observations = []
    for review in reviews:
        source = review.ai_rewritten_remarks or review.remarks
        first_sentence = source.split(". ", 1)[0].strip().rstrip(".")
        observations.append(f"{review.reviewer_name}: {first_sentence}")
    overall_remarks = (
        f"Team consensus from {len(reviews)} "
        f"{'review' if len(reviews) == 1 else 'reviews'}. "
        + "; ".join(observations)
    )[:1800]
    if not overall_remarks.endswith("."):
        overall_remarks += "."
    return TeamReviewsOut(
        reviews=reviews,
        overall_rating=overall_rating,
        overall_remarks=overall_remarks,
        review_count=len(reviews),
    )


@router.get("/links/{link_id}/team-reviews", response_model=TeamReviewsOut)
async def team_reviews(
    link_id: uuid.UUID,
    user: CurrentUser = Depends(require_capability(caps.VIEW_REVIEW_SCREEN)),
    session: AsyncSession = Depends(get_tenant_db),
) -> TeamReviewsOut:
    await _get_link(session, user, link_id)
    return await _team_reviews_out(session, link_id, user.user_id)


@router.put("/links/{link_id}/team-reviews", response_model=TeamReviewsOut)
async def save_team_review(
    link_id: uuid.UUID,
    body: TeamReviewIn,
    user: CurrentUser = Depends(require_capability(caps.DECIDE_PROFILE)),
    session: AsyncSession = Depends(get_tenant_db),
) -> TeamReviewsOut:
    link = await _get_link(session, user, link_id)
    existing = (
        await session.execute(
            select(CandidateTeamReview).where(
                CandidateTeamReview.job_candidate_link_id == link.id,
                CandidateTeamReview.reviewer_user_id == user.user_id,
            )
        )
    ).scalar_one_or_none()
    if existing is None:
        existing = CandidateTeamReview(
            tenant_id=user.tenant_id,
            job_candidate_link_id=link.id,
            reviewer_user_id=user.user_id,
            rating=body.rating,
            remarks=_clean_review_text(body.remarks),
            ai_rewritten_remarks=(
                _clean_review_text(body.ai_rewritten_remarks)
                if body.ai_rewritten_remarks
                else None
            ),
        )
        session.add(existing)
    else:
        existing.rating = body.rating
        existing.remarks = _clean_review_text(body.remarks)
        existing.ai_rewritten_remarks = (
            _clean_review_text(body.ai_rewritten_remarks)
            if body.ai_rewritten_remarks
            else None
        )
        existing.updated_at = datetime.now(timezone.utc)
    await session.flush()
    await audit(
        session,
        tenant_id=user.tenant_id,
        actor_user_id=user.user_id,
        action="candidate_team_review_saved",
        target_type="job_candidate_link",
        target_id=link.id,
        metadata={"rating": body.rating, "used_ai_rewrite": bool(body.ai_rewritten_remarks)},
    )
    return await _team_reviews_out(session, link.id, user.user_id)


@router.post(
    "/links/{link_id}/team-reviews/rewrite",
    response_model=TeamReviewRewriteOut,
)
async def rewrite_team_review(
    link_id: uuid.UUID,
    body: TeamReviewRewriteIn,
    user: CurrentUser = Depends(require_capability(caps.DECIDE_PROFILE)),
    session: AsyncSession = Depends(get_tenant_db),
) -> TeamReviewRewriteOut:
    await _get_link(session, user, link_id)
    fallback = _clean_review_text(body.remarks)
    used_ai = False
    rewritten = fallback
    try:
        raw = await llm_router.invoke_llm(
            "report_synthesis",
            [
                {
                    "role": "system",
                    "content": (
                        "Rewrite a hiring-team observation for clarity and professionalism. "
                        "Preserve every factual claim and the reviewer's meaning. Do not add "
                        "a decision, score, protected characteristic or new evidence. Return "
                        "only the revised remark in 25 to 80 words."
                    ),
                },
                {"role": "user", "content": body.remarks},
            ],
            session=session,
            timeout=8,
            total_budget=12,
        )
        candidate = _clean_review_text(raw.strip().strip("\"'"))
        if 10 <= len(candidate.split()) <= 120:
            rewritten = candidate
            used_ai = True
    except (llm_router.LLMUnavailableError, ValueError):
        pass
    await audit(
        session,
        tenant_id=user.tenant_id,
        actor_user_id=user.user_id,
        action="candidate_team_review_rewritten",
        target_type="job_candidate_link",
        target_id=link_id,
        metadata={"used_ai": used_ai},
    )
    return TeamReviewRewriteOut(rewritten_remarks=rewritten, used_ai=used_ai)


@router.post("/links/{link_id}/status", response_model=StatusOut)
async def update_pipeline_status(
    link_id: uuid.UUID,
    body: StatusIn,
    user: CurrentUser = Depends(require_capability(caps.UPDATE_PIPELINE_STATUS)),
    session: AsyncSession = Depends(get_tenant_db),
) -> StatusOut:
    """Mandatory pipeline status update (FR-8.4). Fresh candidates cannot be
    moved FORWARD until the outreach + all employer verifications are
    submitted or explicitly overridden (FR-5.5)."""
    link = await _get_link(session, user, link_id)
    new_status = PipelineStatus(body.status)

    # PRD v1.0: employer verification is out of scope (§5 non-goal). A candidate
    # applies openly and completes the 40-aspect questionnaire AT application, so
    # the only forward-gate is that the questionnaire is complete — the old
    # VerificationRequest requirement is removed so open applicants aren't blocked.
    if new_status in FORWARD_STATUSES and link.source == LinkSource.fresh:
        profile = await session.get(Profile, link.profile_id) if link.profile_id else None
        if profile is None or profile.aspects_completed_at is None:
            raise HTTPException(
                status_code=409,
                detail="Candidate has not completed the 40-question application yet",
            )

    entry = PipelineStatusEntry(
        tenant_id=user.tenant_id, job_candidate_link_id=link.id,
        status=new_status, remarks=body.remarks, set_by=user.user_id,
    )
    session.add(entry)
    await session.flush()
    await audit(session, tenant_id=user.tenant_id, actor_user_id=user.user_id,
                action="pipeline_status_updated", target_type="job_candidate_link",
                target_id=link.id, metadata={"status": body.status})
    return StatusOut(link_id=link.id, status=entry.status, remarks=entry.remarks,
                     at=entry.at or datetime.now(timezone.utc))


@router.post(
    "/links/{link_id}/interviews",
    response_model=InterviewOut,
    status_code=status.HTTP_201_CREATED,
)
async def schedule_interview(
    link_id: uuid.UUID,
    body: InterviewIn,
    user: CurrentUser = Depends(require_capability(caps.SCHEDULE_INTERVIEWS)),
    session: AsyncSession = Depends(get_tenant_db),
) -> InterviewOut:
    """Interview invite with .ics, sent ONLY from the tenant's verified
    sending domain (FR-8.3 / claude.md rule 5) — never Gmail/Outlook."""
    link = await _get_link(session, user, link_id)
    candidate = await session.get(Candidate, link.candidate_id)
    tenant = await session.get(Tenant, user.tenant_id)
    if candidate is None or tenant is None:
        raise HTTPException(status_code=404, detail="Candidate or tenant missing")

    job = await session.get(Job, link.job_id)
    job_title = job.title if job is not None else "the role"

    ics_uid = f"{uuid.uuid4()}@{tenant.domain}"
    sent_from = f"no-reply@{tenant.domain}"
    interview = Interview(
        tenant_id=user.tenant_id, job_candidate_link_id=link.id,
        scheduled_at=body.scheduled_at, sent_from_email=sent_from,
        ics_uid=ics_uid, notes=body.notes,
    )
    session.add(interview)
    await session.flush()

    celery_app.send_task(
        "pickready.send_email",
        args=[
            str(user.tenant_id), candidate.email, "interview_invite",
            {
                # The three the template actually addresses the candidate with.
                # Without them an unknown placeholder renders as "", so the
                # invitation went out as "Dear ," / "the position at ,".
                "candidate_name": candidate.full_name or "there",
                "job_title": job_title,
                "company_name": tenant.name,
                "scheduled_at": body.scheduled_at.isoformat(),
                "notes": body.notes,
                "ics_uid": ics_uid,
                "interview_id": str(interview.id),
            },
        ],
    )
    await audit(session, tenant_id=user.tenant_id, actor_user_id=user.user_id,
                action="interview_scheduled", target_type="interview",
                target_id=interview.id, metadata={"link_id": str(link.id)})
    return InterviewOut.model_validate(interview)


@router.get("/jobs/{job_id}", response_model=JobLinksOut)
async def list_job_links(
    job_id: uuid.UUID,
    include_archived: bool = Query(default=False),
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=25, ge=1, le=100),
    user: CurrentUser = Depends(require_capability(caps.VIEW_DATABANK)),
    session: AsyncSession = Depends(get_tenant_db),
) -> JobLinksOut:
    """One page of candidate links for a job, with score/tier/current status.

    Joined and paginated: this used to load every link for the job and then
    fetch each candidate individually, so a job with 400 applicants was 401
    queries and a response nobody could render.
    """
    job = await session.get(Job, job_id)
    if job is None or job.tenant_id != user.tenant_id:
        raise HTTPException(status_code=404, detail="Job not found")

    filters = [JobCandidateLink.job_id == job.id]
    if not include_archived:
        filters.append(JobCandidateLink.archived_at.is_(None))

    total = (
        await session.execute(
            select(func.count()).select_from(JobCandidateLink).where(*filters)
        )
    ).scalar_one()
    rows = (
        await session.execute(
            select(JobCandidateLink, Candidate)
            .join(Candidate, Candidate.id == JobCandidateLink.candidate_id)
            .where(*filters)
            # `id` makes the order total, so page boundaries are stable when
            # two links share a score.
            .order_by(
                JobCandidateLink.match_score.desc().nulls_last(),
                JobCandidateLink.id,
            )
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
    ).all()
    latest = await _latest_status(session, [link.id for link, _ in rows])
    out: list[LinkOut] = []
    for link, candidate in rows:
        entry = latest.get(link.id)
        out.append(LinkOut(
            link_id=link.id,
            candidate=CandidateOut.model_validate(candidate),
            profile_id=link.profile_id,
            source=link.source,
            source_type=link.source_type,
            source_type_label=source_type_label(link.source_type),
            tier=link.tier,
            # Numeric parameter scores are internal ranking data and never
            # cross this boundary (claude.md) — comments + scoring_mode only.
            breakdown=client_breakdown(link.match_breakdown_json),
            # Comments-only projection (always present; see services.matching).
            **ranking_payload(link.match_breakdown_json),
            hm_access_granted=link.hm_access_granted,
            archived_at=link.archived_at,
            current_status=entry.status if entry else None,
            status_remarks=entry.remarks if entry else None,
        ))
    total_pages = max(1, (int(total) + page_size - 1) // page_size)
    return JobLinksOut(
        job_id=job.id,
        links=out,
        total=int(total),
        page=page,
        page_size=page_size,
        total_pages=total_pages,
        has_next=page < total_pages,
        has_previous=page > 1,
    )
