"""Deterministic extraction: one file in, one ParsedArtifact out.

No model call anywhere in this module, by design (master brief sections 12 and
31): parsing is the cheap, reproducible layer that runs BEFORE any reasoning
spend, and an outage in a provider must not be an outage in extraction. No
candidate code is ever executed -- source files are read as text and analysed
with structural heuristics, archives are inspected before extraction, and
binary CAD formats get metadata-level readers only.

Every parser is total: a corrupt or hostile file yields an artifact with
`supported=False` and a stated limitation rather than an exception, so one bad
file costs one artifact and never the project (graceful degradation is a
recorded outcome here, not a silent fallback -- the limitation text IS the
record).
"""
from __future__ import annotations

import csv
import io
import json
import re
import struct
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from typing import Any

from app.services.projects import formats
from app.services.projects.formats import Classification, classify
from app.services.projects.limits import ProjectLimits

# ── Output shape ─────────────────────────────────────────────────────────────


@dataclass
class ParsedArtifact:
    path: str
    family: str
    label: str
    supported: bool
    size_bytes: int
    limitation: str | None = None
    #: Deterministic signals keyed by signal name. Values are JSON-safe.
    signals: dict[str, Any] = field(default_factory=dict)
    #: Bounded, normalised text kept for evidence context (documents mainly).
    text_excerpt: str = ""


# ── Technology knowledge tables (data, not branches) ─────────────────────────

#: Dependency name (lowercased) -> technology label surfaced as evidence.
KNOWN_TECHNOLOGIES: dict[str, str] = {
    # Python
    "django": "Django", "flask": "Flask", "fastapi": "FastAPI",
    "sqlalchemy": "SQLAlchemy", "celery": "Celery", "pandas": "pandas",
    "numpy": "NumPy", "scipy": "SciPy", "torch": "PyTorch",
    "tensorflow": "TensorFlow", "scikit-learn": "scikit-learn",
    "pytest": "pytest", "alembic": "Alembic", "pydantic": "Pydantic",
    "matplotlib": "Matplotlib", "opencv-python": "OpenCV",
    # JavaScript / TypeScript
    "react": "React", "next": "Next.js", "vue": "Vue.js", "angular": "Angular",
    "express": "Express", "nestjs": "NestJS", "@nestjs/core": "NestJS",
    "svelte": "Svelte", "jest": "Jest", "vitest": "Vitest",
    "typescript": "TypeScript", "tailwindcss": "Tailwind CSS",
    "electron": "Electron", "three": "Three.js",
    # Data stores / infra
    "postgres": "PostgreSQL", "pg": "PostgreSQL", "psycopg2": "PostgreSQL",
    "asyncpg": "PostgreSQL", "mysql": "MySQL", "mysql2": "MySQL",
    "mongodb": "MongoDB", "mongoose": "MongoDB", "pymongo": "MongoDB",
    "redis": "Redis", "ioredis": "Redis", "kafka-python": "Kafka",
    "kafkajs": "Kafka", "amqplib": "RabbitMQ", "pika": "RabbitMQ",
    "elasticsearch": "Elasticsearch", "sqlite3": "SQLite",
    # Java / JVM
    "spring-boot-starter": "Spring Boot", "spring-boot": "Spring Boot",
    "hibernate-core": "Hibernate", "junit": "JUnit",
    # Other ecosystems
    "rails": "Ruby on Rails", "laravel/framework": "Laravel",
    "gin-gonic/gin": "Gin", "actix-web": "Actix Web", "rocket": "Rocket",
    "boto3": "AWS SDK", "aws-sdk": "AWS SDK", "firebase-admin": "Firebase",
    "stripe": "Stripe", "razorpay": "Razorpay",
}

#: Import/include names that indicate a technology when no manifest names it.
IMPORT_TECHNOLOGIES: dict[str, str] = {
    "django": "Django", "flask": "Flask", "fastapi": "FastAPI",
    "sqlalchemy": "SQLAlchemy", "celery": "Celery", "torch": "PyTorch",
    "tensorflow": "TensorFlow", "sklearn": "scikit-learn", "pandas": "pandas",
    "numpy": "NumPy", "react": "React", "express": "Express",
    "redis": "Redis", "psycopg2": "PostgreSQL", "pymongo": "MongoDB",
    "boto3": "AWS SDK", "cv2": "OpenCV",
}

_ROUTE_PATTERNS = (
    re.compile(r"@(?:app|router|blueprint)\.(?:get|post|put|patch|delete|route)\(", re.I),
    re.compile(r"\b(?:app|router)\.(?:get|post|put|patch|delete)\s*\(\s*['\"]/", re.I),
    re.compile(r"@(?:Get|Post|Put|Patch|Delete|Request)Mapping", re.I),
    re.compile(r"\burlpatterns\s*=", re.I),
)
_AUTH_PATTERN = re.compile(
    r"\b(?:jwt|oauth|passport|authenticate|authorization|login_required|"
    r"bcrypt|argon2|session\s*\[|csrf)\b",
    re.I,
)
_DB_PATTERN = re.compile(
    r"\b(?:select\s+.+\s+from|insert\s+into|create\s+table|findone|"
    r"aggregate\(|objects\.filter|session\.query|prisma\.)",
    re.I,
)
_TEST_PATH = re.compile(r"(?:^|/)(?:tests?|__tests__|spec)(?:/|_|\.)|_test\.|\.test\.|\.spec\.", re.I)
_ERROR_HANDLING = re.compile(r"\b(?:try|except|catch|rescue|recover\(|Result<|panic::)\b")
_CONCURRENCY = re.compile(
    r"\b(?:async\s+def|await|goroutine|go\s+func|threading|multiprocessing|"
    r"tokio|std::thread|CompletableFuture|Promise\.all)\b"
)

_IMPORT_PATTERNS = (
    re.compile(r"^\s*import\s+([\w.]+)", re.M),
    re.compile(r"^\s*from\s+([\w.]+)\s+import", re.M),
    re.compile(r"""require\(\s*['"]([^'"]+)['"]\s*\)"""),
    re.compile(r"""^\s*import\s+.*?from\s+['"]([^'"]+)['"]""", re.M),
    re.compile(r"^\s*#include\s+[<\"]([\w./]+)[>\"]", re.M),
    re.compile(r"^\s*using\s+([\w.]+);", re.M),
)

_DEFINITION_PATTERN = re.compile(
    r"^\s*(?:def |fn |func |function |public |private |protected |static |"
    r"module |always @|entity |architecture )",
    re.M,
)
_CLASS_PATTERN = re.compile(
    r"^\s*(?:class |interface |struct |trait |impl |component )", re.M
)


def _decode(data: bytes, limit: int) -> str:
    text = data[: limit * 4].decode("utf-8", errors="replace")
    return text[:limit]


def _relevant_imports(text: str) -> list[str]:
    found: list[str] = []
    for pattern in _IMPORT_PATTERNS:
        for match in pattern.findall(text):
            root = str(match).split("/")[0].split(".")[0].strip().lower()
            if root and root not in found:
                found.append(root)
    return found[:60]


# ── Per-family parsers ───────────────────────────────────────────────────────


def _parse_source(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    text = _decode(data, limits.max_text_chars_per_file)
    imports = _relevant_imports(text)
    technologies = sorted(
        {IMPORT_TECHNOLOGIES[name] for name in imports if name in IMPORT_TECHNOLOGIES}
    )
    signals: dict[str, Any] = {
        "language": cls.label,
        "line_count": text.count("\n") + 1,
        "imports": imports,
        "technologies": technologies,
        "definition_count": len(_DEFINITION_PATTERN.findall(text)),
        "class_count": len(_CLASS_PATTERN.findall(text)),
        "is_test": bool(_TEST_PATH.search(path)),
        "has_routes": any(p.search(text) for p in _ROUTE_PATTERNS),
        "has_auth": bool(_AUTH_PATTERN.search(text)),
        "has_db_access": bool(_DB_PATTERN.search(text)),
        "has_error_handling": bool(_ERROR_HANDLING.search(text)),
        "has_concurrency": bool(_CONCURRENCY.search(text)),
    }
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label, supported=True,
        size_bytes=len(data), signals=signals,
    )


def _parse_manifest(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    text = _decode(data, limits.max_text_chars_per_file)
    lowered_name = path.replace("\\", "/").rsplit("/", 1)[-1].lower()
    dependencies: list[str] = []
    if lowered_name == "package.json":
        try:
            payload = json.loads(text)
            for key in ("dependencies", "devDependencies"):
                dependencies.extend((payload.get(key) or {}).keys())
        except (json.JSONDecodeError, AttributeError):
            return ParsedArtifact(
                path=path, family=cls.family, label=cls.label, supported=False,
                size_bytes=len(data),
                limitation="The manifest could not be parsed as JSON.",
            )
    elif lowered_name == "requirements.txt":
        for line in text.splitlines():
            name = re.split(r"[<>=!\[;#]", line.strip(), 1)[0].strip()
            if name and not name.startswith(("-", "#")):
                dependencies.append(name)
    elif lowered_name in {"pyproject.toml", "cargo.toml"}:
        dependencies.extend(re.findall(r'^\s*([\w][\w.-]*)\s*=\s*["{\[]', text, re.M))
    elif lowered_name == "go.mod":
        dependencies.extend(re.findall(r"^\s*([\w./-]+)\s+v[\d.]", text, re.M))
    elif lowered_name in {"pom.xml", "build.gradle", "build.gradle.kts"}:
        dependencies.extend(re.findall(r"<artifactId>([\w.-]+)</artifactId>", text))
        dependencies.extend(
            re.findall(r"""implementation\s*[('"]+([\w.:-]+)""", text)
        )
    technologies = sorted(
        {
            KNOWN_TECHNOLOGIES[dep.lower()]
            for dep in dependencies
            if dep.lower() in KNOWN_TECHNOLOGIES
        }
    )
    is_docker = lowered_name.startswith("dockerfile") or lowered_name.startswith(
        "docker-compose"
    )
    if is_docker:
        technologies = sorted(set(technologies) | {"Docker"})
    signals: dict[str, Any] = {
        "manifest": lowered_name,
        "dependency_count": len(dependencies),
        "dependencies": dependencies[:80],
        "technologies": technologies,
        "is_containerised": is_docker,
    }
    if lowered_name.startswith("docker-compose"):
        signals["compose_services"] = re.findall(
            r"^\s{2}([\w-]+):\s*$", text, re.M
        )[:20]
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label, supported=True,
        size_bytes=len(data), signals=signals,
    )


def _parse_document(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    label = cls.label.lower()
    text = ""
    page_count = None
    limitation = None
    if label == "pdf":
        try:
            from pypdf import PdfReader  # noqa: PLC0415 -- heavy import, parse path only

            reader = PdfReader(io.BytesIO(data))
            page_count = len(reader.pages)
            pages = [(page.extract_text() or "") for page in reader.pages[:40]]
            text = "\n".join(pages)
        except Exception:  # noqa: BLE001 -- corrupt/encrypted PDF is a recorded limitation
            return ParsedArtifact(
                path=path, family=cls.family, label=cls.label, supported=False,
                size_bytes=len(data),
                limitation="The PDF could not be read; it may be corrupt, "
                "encrypted, or image-only.",
            )
    elif label == "docx":
        try:
            import docx  # noqa: PLC0415

            document = docx.Document(io.BytesIO(data))
            parts = [p.text for p in document.paragraphs if p.text]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        except Exception:  # noqa: BLE001
            return ParsedArtifact(
                path=path, family=cls.family, label=cls.label, supported=False,
                size_bytes=len(data),
                limitation="The DOCX file could not be read as a document.",
            )
    else:
        text = _decode(data, limits.max_text_chars_per_file)
        if label in {"html", "htm"}:
            text = re.sub(r"<[^>]+>", " ", text)
    text = text[: limits.max_text_chars_per_file]
    headings = re.findall(r"^#{1,4}\s+(.{2,120})$", text, re.M)[:20]
    signals: dict[str, Any] = {
        "word_count": len(re.findall(r"\b[\w'-]+\b", text)),
        "headings": [h.strip() for h in headings],
        "is_readme": "readme" in path.lower(),
    }
    if page_count is not None:
        signals["page_count"] = page_count
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label, supported=True,
        size_bytes=len(data), limitation=limitation, signals=signals,
        text_excerpt=text[:4000],
    )


def _parse_data(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    text = _decode(data, limits.max_text_chars_per_file)
    normalised = path.replace("\\", "/").lower()
    signals: dict[str, Any] = {
        "is_ci_config": (
            "/.github/workflows/" in normalised
            or normalised.endswith((".gitlab-ci.yml", "azure-pipelines.yml"))
            or "/.circleci/" in normalised
        ),
        "is_infrastructure": normalised.endswith((".tf", ".tfvars"))
        or "kubernetes" in normalised
        or bool(re.search(r"^\s*(?:apiVersion|kind):", text, re.M)),
        "key_count": len(re.findall(r"^\s*[\w.-]+\s*[:=]", text, re.M)),
    }
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label, supported=True,
        size_bytes=len(data), signals=signals,
    )


def _parse_spreadsheet(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    label = cls.label.lower()
    signals: dict[str, Any] = {}
    if label in {"csv", "tsv"}:
        text = _decode(data, limits.max_text_chars_per_file)
        delimiter = "\t" if label == "tsv" else ","
        try:
            rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
        except csv.Error:
            rows = []
        signals["row_count"] = max(len(rows) - 1, 0)
        signals["columns"] = [c.strip()[:60] for c in (rows[0] if rows else [])][:30]
    elif label == "xlsx":
        # An XLSX is a ZIP of XML; sheet names are deterministic metadata that
        # needs no spreadsheet library.
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                workbook = archive.read("xl/workbook.xml").decode(
                    "utf-8", errors="replace"
                )
            signals["sheet_names"] = re.findall(r'<sheet[^>]*name="([^"]+)"', workbook)[:20]
        except (zipfile.BadZipFile, KeyError, OSError):
            return ParsedArtifact(
                path=path, family=cls.family, label=cls.label, supported=False,
                size_bytes=len(data),
                limitation="The spreadsheet could not be read as an XLSX file.",
            )
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label, supported=True,
        size_bytes=len(data), signals=signals,
    )


def _parse_notebook(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    try:
        payload = json.loads(_decode(data, limits.max_text_chars_per_file * 4))
        cells = payload.get("cells") or []
    except (json.JSONDecodeError, AttributeError):
        return ParsedArtifact(
            path=path, family=cls.family, label=cls.label, supported=False,
            size_bytes=len(data),
            limitation="The notebook could not be parsed as JSON.",
        )
    code = "\n".join(
        "".join(cell.get("source") or [])
        for cell in cells
        if isinstance(cell, dict) and cell.get("cell_type") == "code"
    )
    imports = _relevant_imports(code)
    signals: dict[str, Any] = {
        "cell_count": len(cells),
        "code_cell_count": sum(
            1 for c in cells if isinstance(c, dict) and c.get("cell_type") == "code"
        ),
        "imports": imports,
        "technologies": sorted(
            {IMPORT_TECHNOLOGIES[n] for n in imports if n in IMPORT_TECHNOLOGIES}
        ),
        "language": (payload.get("metadata") or {})
        .get("kernelspec", {})
        .get("language", "python"),
    }
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label, supported=True,
        size_bytes=len(data), signals=signals,
    )


# ── CAD / engineering metadata readers ───────────────────────────────────────


def _parse_cad(path: str, data: bytes, cls: Classification, limits: ProjectLimits) -> ParsedArtifact:
    label = cls.label.lower()
    signals: dict[str, Any] = {"format": cls.label}
    limitation = None
    if label in {"step", "stp", "iges", "igs"}:
        text = _decode(data, limits.max_text_chars_per_file)
        description = re.search(
            r"FILE_DESCRIPTION\s*\(\s*\(\s*'([^']*)'", text
        )
        schema = re.search(r"FILE_SCHEMA\s*\(\s*\(\s*'([^']*)'", text)
        if description:
            signals["file_description"] = description.group(1)[:200]
        if schema:
            signals["schema"] = schema.group(1)[:80]
        signals["product_count"] = len(re.findall(r"=\s*PRODUCT\s*\(", text))
        signals["assembly_relationships"] = len(
            re.findall(r"NEXT_ASSEMBLY_USAGE_OCCURRENCE", text)
        )
    elif label == "stl":
        header = data[:512]
        if header.lstrip().lower().startswith(b"solid"):
            text = _decode(data, limits.max_text_chars_per_file)
            signals["facet_count"] = text.count("facet normal")
            name = re.match(r"\s*solid\s+(.{0,80})", text)
            if name:
                signals["solid_name"] = name.group(1).strip()
        elif len(data) >= 84:
            signals["facet_count"] = struct.unpack("<I", data[80:84])[0]
    elif label == "dxf":
        text = _decode(data, limits.max_text_chars_per_file)
        lines = text.splitlines()
        entities = Counter(
            lines[i + 1].strip()
            for i, line in enumerate(lines[:-1])
            if line.strip() == "0" and lines[i + 1].strip().isalpha()
        )
        signals["entity_counts"] = dict(entities.most_common(15))
        signals["layer_count"] = entities.get("LAYER", 0)
    elif label == "ifc":
        text = _decode(data, limits.max_text_chars_per_file * 4)
        entities = Counter(re.findall(r"=\s*(IFC[A-Z]+)\s*\(", text))
        signals["entity_counts"] = dict(entities.most_common(25))
        schema = re.search(r"FILE_SCHEMA\s*\(\s*\(\s*'([^']*)'", text)
        if schema:
            signals["schema"] = schema.group(1)[:40]
    if not any(value for key, value in signals.items() if key != "format"):
        limitation = (
            "The file was recognised as a CAD artifact but no structured "
            "metadata could be extracted from it."
        )
    return ParsedArtifact(
        path=path, family=cls.family, label=cls.label,
        supported=limitation is None, size_bytes=len(data),
        limitation=limitation, signals=signals,
    )


# ── The router ───────────────────────────────────────────────────────────────

_PARSERS = {
    formats.FAMILY_SOURCE: _parse_source,
    formats.FAMILY_MANIFEST: _parse_manifest,
    formats.FAMILY_DOCUMENT: _parse_document,
    formats.FAMILY_DATA: _parse_data,
    formats.FAMILY_SPREADSHEET: _parse_spreadsheet,
    formats.FAMILY_NOTEBOOK: _parse_notebook,
    formats.FAMILY_CAD: _parse_cad,
}


def parse_file(path: str, data: bytes, limits: ProjectLimits) -> ParsedArtifact:
    """Route one file to its parser. Total: unsupported formats come back as
    recorded artifacts, never as exceptions."""
    cls = classify(path)
    if not cls.supported or cls.family not in _PARSERS:
        return ParsedArtifact(
            path=path, family=cls.family, label=cls.label, supported=False,
            size_bytes=len(data), limitation=cls.limitation,
        )
    return _PARSERS[cls.family](path, data, cls, limits)
