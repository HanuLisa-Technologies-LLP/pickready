"""Resume parsing pipeline (ESD §9, FR-6.2).

Raw PDF/DOCX -> text extraction (pypdf / python-docx) -> LLM extraction
(`extraction` task type, the extraction tier -- narrow, mechanical, must not evaluate)
into a fixed structured schema -> profiles.parsed_fields_json. Also sets
profiles.resume_text and the voyage-context-4 embedding used by the semantic
matching stage.

THIS IS ALSO WHERE YUKTI'S PRE-SCREEN GRADE IS PRODUCED (spec-doc6 §4.4).
`parse_resume` ends by calling `hiring.prescreen.grade_profile`, which writes an
A / B / C / Hold onto every application this profile is attached to. It is here
and not on each upload route on purpose: five routes accept a resume (a
candidate applying, a candidate replacing their main resume, the databank bulk
upload, the provider-side upload, the seeding path) and all five already enqueue
`pickready.parse_resume`, so hanging the grade off the parse means no route had
to learn about grading and no future route can forget to.

It runs even when nothing could be extracted. A scanned image resume produces
`Hold`, which is a graded outcome meaning a person should look, and that is a
better dashboard cell than an empty one: an empty cell is indistinguishable from
a candidate nobody has got to yet.
"""
from __future__ import annotations

import io
import json
import logging
import re
import uuid
from typing import Any

from sqlalchemy.ext.asyncio import AsyncSession

from app.models import Profile
from app.services import llm_router
from app.services.embeddings import embed
from app.services.resume_storage import ResumeStorageError, fetch_resume_bytes, profile_has_resume
from app.prompts import registry

logger = logging.getLogger(__name__)

_DOWNLOAD_TIMEOUT = 60.0

#: An empty parsed-fields document in the fixed schema (used when a resume has
#: no extractable text or the LLM extraction can't be parsed  -  the profile is
#: still stored so nothing crashes the Celery task).
_EMPTY_PARSED_FIELDS: dict[str, Any] = {
    "skills": [],
    "total_experience_years": None,
    "education": [],
    "employment_history": [],
}

#: The fixed extraction schema (ESD §9): keep in sync with the prompt below.
PARSED_FIELDS_SCHEMA: dict[str, Any] = {
    "skills": [],
    "total_experience_years": None,
    "education": [],
    "employment_history": [],  # [{company, title, start, end}]
}


class ResumeParsingError(RuntimeError):
    pass


# ── Text extraction ──────────────────────────────────────────────────────────
#
# Extraction is defensive: a corrupt, empty, image-only, or wrong-format file
# yields "" rather than raising, so the Celery `parse_resume` task never
# crash-loops on unparseable content (a genuinely transient failure  -  e.g. the
# Cloudinary download  -  still propagates from `parse_resume` and is retried).


def extract_text(data: bytes, filename: str) -> str:
    """Extract plain text from a PDF or DOCX resume.

    Returns clean text, or "" when nothing extractable is found (empty file,
    corrupt archive, scanned/image-only PDF, unknown format). Never raises on
    content problems  -  the caller decides how to handle an empty result.
    """
    if not data:
        return ""
    lowered = (filename or "").lower()
    if lowered.endswith(".pdf"):
        return _extract_pdf(data)
    if lowered.endswith(".docx"):
        return _extract_docx(data)
    # ASSUMPTION: unknown extensions (Cloudinary raw URLs may drop the original
    # extension)  -  sniff the magic bytes: PDFs start with "%PDF", DOCX is a ZIP
    # ("PK"). Fall back to trying both.
    if data[:4] == b"%PDF":
        return _extract_pdf(data)
    if data[:2] == b"PK":
        return _extract_docx(data)
    return _extract_pdf(data) or _extract_docx(data)


def _extract_pdf(data: bytes) -> str:
    """Best-effort PDF text extraction; "" on empty/corrupt/scanned input."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:  # noqa: BLE001  -  corrupt/encrypted PDF, not a crash
        logger.warning("resume_parsing.pdf_extract_failed error=%s", type(exc).__name__)
        return ""
    return "\n".join(pages).strip()


def _extract_docx(data: bytes) -> str:
    """Best-effort DOCX text extraction (paragraphs + tables); "" on failure."""
    try:
        import docx  # python-docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells if cell.text)
    except Exception as exc:  # noqa: BLE001  -  not a real .docx / corrupt archive
        logger.warning("resume_parsing.docx_extract_failed error=%s", type(exc).__name__)
        return ""
    return "\n".join(parts).strip()


# ── Cheap contact identity (databank bulk upload, 2026-07-28) ────────────────
#
# The databank uploader has to create or find a candidate row BEFORE it can
# create the job link, and it must not wait for the real parse to happen (that
# is a Celery task, claude.md rule 4, and it calls an LLM). So it needs one
# cheap, local, deterministic answer to "whose resume is this".
#
# These are pure regex/heuristic functions over already-extracted text. They do
# NO network I/O and NO LLM call, and they never replace `parse_resume`, which
# still runs asynchronously afterwards and fills in skills, experience,
# education, employment history and the embedding.

_EMAIL_RE = re.compile(
    r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,24}"
)
# Deliberately loose: resumes write phone numbers every way imaginable.
_PHONE_RE = re.compile(r"(?:\+?\d[\d\s\-().]{7,17}\d)")

#: Words that mean a line is a section heading, not somebody's name.
_NON_NAME_TOKENS = frozenset(
    {
        "resume", "curriculum", "vitae", "cv", "profile", "summary", "objective",
        "contact", "address", "phone", "email", "experience", "education",
        "skills", "projects", "certifications", "linkedin", "github",
    }
)


def extract_email(text: str) -> str | None:
    """First plausible email address in the text, lowercased."""
    if not text:
        return None
    match = _EMAIL_RE.search(text)
    return match.group(0).strip().lower() if match else None


def extract_phone(text: str) -> str | None:
    """First plausible phone number, digits and a leading + only.

    Truncated to the column width (20) rather than dropped, because a partially
    captured number is still a lead a recruiter can act on.
    """
    if not text:
        return None
    match = _PHONE_RE.search(text)
    if not match:
        return None
    cleaned = re.sub(r"[^\d+]", "", match.group(0))
    return cleaned[:20] or None


def extract_full_name(text: str) -> str | None:
    """Best-effort candidate name from the first few lines of a resume.

    Heuristic and openly so: resumes overwhelmingly put the person's name on
    the first non-empty line. A line is rejected when it is too long, contains
    an email or a digit, or reads as a section heading. Returns None rather
    than a guess it does not believe, and the caller falls back to a clearly
    marked placeholder identity.
    """
    if not text:
        return None
    for raw in text.splitlines()[:12]:
        line = re.sub(r"\s+", " ", raw).strip(" \t|,-·•")
        if not (2 <= len(line) <= 60):
            continue
        if "@" in line or any(ch.isdigit() for ch in line):
            continue
        words = line.split(" ")
        if not (1 < len(words) <= 5):
            continue
        if any(word.lower().strip(":") in _NON_NAME_TOKENS for word in words):
            continue
        if not all(word[0].isalpha() for word in words if word):
            continue
        return line
    return None


def extract_contact_identity(data: bytes, filename: str) -> dict[str, str | None]:
    """`{full_name, email, phone}` from raw resume bytes. Never raises."""
    try:
        text = extract_text(data, filename)
    except Exception as exc:  # noqa: BLE001  -  identity is best effort
        logger.warning(
            "resume_parsing.identity_extract_failed error=%s", type(exc).__name__
        )
        return {"full_name": None, "email": None, "phone": None}
    return {
        "full_name": extract_full_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
    }


# ── LLM structured extraction ────────────────────────────────────────────────

#: Text in `app/prompts/resume_extraction_system.txt`, loaded through the registry so a
#: wording change is a versioned diff in a prompt file rather than a string
#: literal in a module of code. What is sent is unchanged.
_EXTRACTION_SYSTEM = registry.render("resume_extraction_system")


async def extract_structured_fields(
    resume_text: str, session: AsyncSession | None = None
) -> dict:
    """Run the LLM extraction chain over resume text -> parsed_fields dict.

    Degrades gracefully: empty text short-circuits to the empty schema without
    an LLM call, and unparseable LLM output logs a warning and returns the
    empty schema rather than raising  -  the raw resume text stays as the
    matcher's fallback signal. `llm_router.LLMUnavailableError` (whole provider
    chain down) still propagates so the Celery task's retry policy handles it.
    """
    if not resume_text or not resume_text.strip():
        return dict(_EMPTY_PARSED_FIELDS)
    raw = await llm_router.chat_completion(
        "extraction",
        [
            {"role": "system", "content": _EXTRACTION_SYSTEM},
            {"role": "user", "content": resume_text[:24000]},
        ],
        response_format_json=True,
        session=session,
    )
    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        logger.warning("resume_parsing.extraction_non_json  -  storing empty parsed fields")
        return dict(_EMPTY_PARSED_FIELDS)
    if not isinstance(parsed, dict):
        logger.warning("resume_parsing.extraction_not_object  -  storing empty parsed fields")
        return dict(_EMPTY_PARSED_FIELDS)
    # Normalize to the fixed schema  -  never store surprise keys.
    return {
        "skills": parsed.get("skills") or [],
        "total_experience_years": parsed.get("total_experience_years"),
        "education": parsed.get("education") or [],
        "employment_history": parsed.get("employment_history") or [],
    }


# ── Full pipeline for one profile ────────────────────────────────────────────

async def parse_resume(session: AsyncSession, profile_id: uuid.UUID | str) -> None:
    """Extract text (if needed), run LLM extraction, store parsed fields,
    resume_text, and the voyage-context-4 embedding on the profile."""
    profile = await session.get(Profile, uuid.UUID(str(profile_id)))
    if profile is None:
        raise ValueError(f"Profile {profile_id} not found")

    resume_text = profile.resume_text
    if not resume_text:
        if not profile_has_resume(profile):
            raise ResumeParsingError(
                f"Profile {profile_id} has no complete Cloudinary resume metadata"
            )
        # Download failures (network/5xx) propagate so the task retries  -  only
        # *content* problems (below) are swallowed.
        try:
            data = await fetch_resume_bytes(profile)
        except ResumeStorageError as exc:
            raise ResumeParsingError(str(exc)) from exc
        resume_text = extract_text(data, profile.resume_original_filename)

    if not resume_text or not resume_text.strip():
        # Empty/garbage/scanned resume  -  persist an empty profile and stop.
        # Leaving embedding NULL excludes it from the semantic matching pool
        # (WHERE embedding IS NOT NULL) rather than polluting it with noise.
        logger.warning(
            "resume_parsing.no_extractable_text profile_id=%s, stored empty parsed fields",
            profile_id,
        )
        profile.resume_text = ""
        profile.parsed_fields_json = dict(_EMPTY_PARSED_FIELDS)
        await _prescreen(session, profile)
        await session.commit()
        return

    try:
        parsed = await extract_structured_fields(resume_text, session=session)
    except llm_router.LLMUnavailableError:
        # Indexing the resume must not depend on optional structured extraction.
        # Preserve the raw text + embedding so keyword/semantic matching works,
        # and let a later parse retry enrich the structured fields.
        logger.warning(
            "resume_parsing.extraction_unavailable profile_id=%s, "
            "storing text-only fallback",
            profile_id,
        )
        parsed = dict(_EMPTY_PARSED_FIELDS)

    profile.resume_text = resume_text
    profile.parsed_fields_json = parsed
    profile.embedding = (await embed([resume_text]))[0]
    # Before the commit, so the grade and the parse it was written from land in
    # ONE transaction. A grade committed separately could survive a rolled-back
    # parse and would then describe a resume the row no longer holds, which is
    # the same shape of defect as a `framework_generated_at` stamp with no
    # competency rows behind it.
    await _prescreen(session, profile)
    await session.commit()


async def _prescreen(session: AsyncSession, profile: Profile) -> None:
    """Grade this resume against every job the candidate is linked to.

    Never fatal to the parse. Parsing a resume is what makes the candidate
    searchable, matchable and assessable; a grading failure costs one dashboard
    cell, and taking the parse down with it would cost the candidate their whole
    presence in the product. The failure is logged with the profile it happened
    on rather than swallowed, and the cell stays NULL, which the dashboard
    renders as "not pre-screened" and never as a grade.
    """
    from app.services.hiring import prescreen  # noqa: PLC0415

    # `getattr` on the identifier and not `profile.id`: this whole function is
    # the path that must not raise, and a log line that reaches for an attribute
    # is one AttributeError away from turning a grading failure into a parsing
    # failure, which is the exact outcome the guard exists to prevent.
    profile_id = getattr(profile, "id", None)
    try:
        graded = await prescreen.grade_profile(session, profile)
    except Exception:  # noqa: BLE001 - logged with its subject, never silent
        logger.warning(
            "resume_parsing.prescreen_failed profile_id=%s", profile_id, exc_info=True
        )
        return
    logger.info(
        "resume_parsing.prescreened profile_id=%s applications=%d", profile_id, graded
    )
