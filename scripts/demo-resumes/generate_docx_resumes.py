import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

emails = [
    "manjuhr.m@gmail.com",
    "manjuhr@outlook.com",
    "saravankumar2503@gmail.com",
    "sarkar120806@gmail.com",
    "kvsr101112@gmail.com",
    "saravankumarmk@gmail.com",
    "126004238@sastra.ac.in"
]

phones = [
    "+91 9398687358",
    "+91 9652802233"
]

output_dir = r"C:\Users\Saravan Kumar\OneDrive\Desktop\generalresumes"

resumes_data = [
    # 01 - Java Backend
    {
        "num": "01",
        "name": "Akash Rao",
        "role": "Senior Java Backend Developer",
        "summary": "Senior Java Backend Developer with 4.5+ years of experience designing and deploying high-concurrency microservices, event-driven architectures, and distributed systems. Expert in Java 17, Spring Boot, Apache Kafka, and PostgreSQL, with a proven track record of reducing API latency by 40% and maintaining 99.99% uptime for enterprise-scale platforms.",
        "skills": [
            ("Languages", "Java 17/11, SQL, Go, Bash"),
            ("Frameworks & Libraries", "Spring Boot, Spring Cloud, Hibernate, Spring Security, JUnit 5, Mockito"),
            ("Architecture & Patterns", "Microservices, RESTful APIs, Event-Driven Architecture, CQRS, Domain-Driven Design"),
            ("Databases & Caching", "PostgreSQL, MySQL, Redis, Hazelcast"),
            ("Messaging & Integration", "Apache Kafka, RabbitMQ, gRPC"),
            ("Cloud & DevOps", "Docker, Kubernetes, AWS (EC2, S3, RDS), Jenkins, Maven, Git")
        ],
        "experience": [
            {
                "title": "Senior Software Engineer",
                "company": "TechMetrics Solutions, Bangalore, India",
                "period": "Oct 2022 – Present",
                "bullets": [
                    "Architected and implemented 12+ scalable microservices using Java 17 and Spring Boot, handling over 15 million daily API transactions.",
                    "Transitioned synchronous REST communication to asynchronous event streams with Apache Kafka, improving throughput by 65% during peak traffic.",
                    "Configured Redis distributed caching strategies for session management and database query caching, cutting database load by 45%.",
                    "Led a team of 4 junior developers, establishing code review guidelines, CI/CD pipelines in Jenkins, and unit testing coverage standards (85%+ coverage)."
                ]
            },
            {
                "title": "Software Engineer",
                "company": "Infosys Technologies, Mysore, India",
                "period": "Jul 2019 – Sep 2022",
                "bullets": [
                    "Developed core backend services for a global banking client utilizing Spring Boot, Hibernate, and Oracle Database.",
                    "Optimized legacy database queries and database indexing, reducing p99 latency from 1.2s to 280ms for financial reporting endpoints.",
                    "Implemented OAuth2 and JWT-based authentication mechanisms with Spring Security, securing internal microservices endpoints.",
                    "Managed container deployment on AWS EKS using Docker and Kubernetes helm charts."
                ]
            }
        ],
        "projects": [
            {
                "name": "High-Throughput Order Management System",
                "tech": "Java 17, Spring Boot, Kafka, Redis, PostgreSQL",
                "bullets": [
                    "Engineered a resilient order processing engine designed to handle up to 10,000 orders per second using Kafka event partitions and optimistic locking in PostgreSQL.",
                    "Integrated Redis for real-time inventory locking, mitigating race conditions during high-demand flash sale events."
                ]
            },
            {
                "name": "Enterprise Audit Logging Microservice",
                "tech": "Java 11, Spring Boot, Elasticsearch, Logstash",
                "bullets": [
                    "Built a centralized audit system aggregating system logs across 20+ microservices with sub-second search and structured parsing capabilities."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science and Engineering | VIT University, Vellore, India | Graduated: May 2019 | CGPA: 8.7/10.0"
        ],
        "certifications": [
            "Oracle Certified Professional: Java SE 11 Developer",
            "AWS Certified Developer – Associate"
        ],
        "achievements": [
            "Received 'Best Technical Innovator Award' at TechMetrics Solutions for leading the Kafka event stream migration."
        ]
    },
    # 02 - Java Backend
    {
        "num": "02",
        "name": "Arjun Kumar",
        "role": "Java Backend Engineer",
        "summary": "Backend Software Engineer with 2.5 years of experience specializing in Java, Spring Boot, and relational database management. Skilled in building RESTful microservices, containerizing workloads, and implementing automated testing frameworks. Adept at collaborating in Agile environments to deliver robust enterprise solutions.",
        "skills": [
            ("Languages", "Java 11/8, SQL, HTML/CSS"),
            ("Backend Technologies", "Spring Boot, Spring Data JPA, Hibernate, REST APIs, Liquibase"),
            ("Databases & Storage", "PostgreSQL, MySQL, H2"),
            ("Testing & Tools", "JUnit, Mockito, Postman, Git, Docker, Maven, Gradle, JIRA"),
            ("Concepts", "Object-Oriented Design, SOLID Principles, Data Structures & Algorithms")
        ],
        "experience": [
            {
                "title": "Software Developer",
                "company": "NexaWare Technologies, Hyderabad, India",
                "period": "Jan 2022 – Present",
                "bullets": [
                    "Designed and maintained RESTful web services in Spring Boot for an e-commerce fulfillment platform processing 50,000+ daily orders.",
                    "Optimized JPA and Hibernate entities using lazy loading and join fetch queries, eliminating N+1 query problems and accelerating database response times by 35%.",
                    "Built comprehensive unit and integration test suites using JUnit 5 and Mockito, raising codebase coverage from 60% to 88%.",
                    "Participated in bi-weekly Agile sprints, participating in sprint planning, backlog grooming, and system design discussions."
                ]
            },
            {
                "title": "Java Developer Intern",
                "company": "Wipro Digital, Hyderabad, India",
                "period": "Jun 2021 – Dec 2021",
                "bullets": [
                    "Assisted in developing REST APIs for an internal employee onboarding portal using Java 8 and Spring Boot.",
                    "Implemented automated database migration scripts using Liquibase, ensuring seamless schema updates across staging and production environments."
                ]
            }
        ],
        "projects": [
            {
                "name": "Smart Restaurant Ordering & POS API",
                "tech": "Java, Spring Boot, PostgreSQL, Docker",
                "bullets": [
                    "Developed a complete REST API suite for restaurant table reservation, live order tracking, and billing integration.",
                    "Configured Docker Compose setup for local environment orchestration involving Spring Boot services and PostgreSQL containers."
                ]
            },
            {
                "name": "Personal Finance Tracker API",
                "tech": "Java, Spring Boot, MySQL, JWT",
                "bullets": [
                    "Created a secure backend API allowing users to log expenses, categorize transactions, and generate monthly CSV financial reports."
                ]
            }
        ],
        "education": [
            "B.Tech in Information Technology | JNTU Hyderabad, India | Graduated: Jun 2021 | Percentage: 82%"
        ],
        "certifications": [
            "Java SE 11 Programmer I (1Z0-815)"
        ],
        "achievements": [
            "Secured 1st Place in internal company Hackathon for developing a prototype inventory forecasting tool."
        ]
    },
    # 03 - Java Backend
    {
        "num": "03",
        "name": "Dinesh Bose",
        "role": "Associate Java Developer",
        "summary": "Junior Java Backend Engineer with 1 year of hands-on software development experience building scalable API web applications using Java 17, Spring Boot, and MySQL. Eager to leverage strong computer science fundamentals, data structures knowledge, and modern Java features to build reliable cloud-native services.",
        "skills": [
            ("Languages", "Java 17, Python (Basic), SQL"),
            ("Frameworks", "Spring Boot, Spring MVC, Hibernate/JPA"),
            ("Web & APIs", "RESTful APIs, Swagger/OpenAPI, JSON, HTTP"),
            ("Databases", "MySQL, PostgreSQL"),
            ("Tools & Methodologies", "Git, Maven, Docker, IntelliJ IDEA, Postman, Agile/Scrum")
        ],
        "experience": [
            {
                "title": "Associate Java Developer",
                "company": "TechMahindra, Pune, India",
                "period": "Jul 2023 – Present",
                "bullets": [
                    "Implemented 15+ backend endpoints using Spring Boot and Spring MVC for a logistics management system.",
                    "Standardized API documentation by integrating OpenAPI/Swagger, streamlining front-end integration workflows.",
                    "Refactored legacy monolithic backend modules into modern modular Java packages, improving maintainability and reducing build times by 20%.",
                    "Debugged production issues using log analysis tools and fixed critical bugs in order processing workflows."
                ]
            }
        ],
        "projects": [
            {
                "name": "Hospital Management & Bed Booking System",
                "tech": "Java 17, Spring Boot, MySQL, Thymeleaf",
                "bullets": [
                    "Built a web platform facilitating patient registration, appointment scheduling, and real-time hospital bed availability tracking.",
                    "Secured user authentication using Spring Security with role-based access control (Admin, Doctor, Patient)."
                ]
            },
            {
                "name": "Automated Inventory Alerting Service",
                "tech": "Java, Spring Boot, MailSender, MySQL",
                "bullets": [
                    "Developed a microservice that runs scheduled tasks to monitor warehouse inventory thresholds and send automated email alerts to suppliers."
                ]
            }
        ],
        "education": [
            "B.E. in Computer Science | Pune University, Pune, India | Graduated: May 2023 | CGPA: 8.1/10.0"
        ],
        "certifications": [
            "Udemy Certified Spring Boot & Hibernate Developer"
        ],
        "achievements": [
            "Published a technical paper on Distributed Database Indexing in College Research Journal."
        ]
    },

    # 04 - Python Backend
    {
        "num": "04",
        "name": "Priya Sen",
        "role": "Senior Python Backend Developer",
        "summary": "Python Backend Engineer with 3.5 years of experience architecting high-performance asynchronous web applications, background worker queues, and REST/gRPC APIs using Python, FastAPI, Django, PostgreSQL, and Redis. Expertise in building data-intensive microservices and deploying serverless applications on AWS.",
        "skills": [
            ("Languages", "Python 3.11, SQL, Bash"),
            ("Frameworks", "FastAPI, Django, Django REST Framework, Flask, Celery"),
            ("Databases & Storage", "PostgreSQL, Redis, MongoDB, SQLAlchemy"),
            ("Cloud & Infrastructure", "AWS (Lambda, ECS, S3, RDS), Docker, Terraform, GitHub Actions"),
            ("Architecture", "Asynchronous Programming (asyncio), Microservices, REST, WebSockets")
        ],
        "experience": [
            {
                "title": "Backend Developer",
                "company": "CloudScale Systems, Bangalore, India",
                "period": "Mar 2022 – Present",
                "bullets": [
                    "Designed and deployed asynchronous microservices using FastAPI and asyncio, achieving 4,500+ requests per second with sub-50ms average latency.",
                    "Implemented background task queues with Celery and Redis to handle asynchronous invoice generation and bulk notification processing.",
                    "Migrated legacy Django monolithic application to serverless architecture using AWS Lambda and API Gateway, cutting infrastructure costs by 30%.",
                    "Integrated Pytest and GitHub Actions, establishing automated testing pipelines that maintain 90% code coverage."
                ]
            },
            {
                "title": "Python Developer",
                "company": "Mindtree, Bangalore, India",
                "period": "Nov 2020 – Feb 2022",
                "bullets": [
                    "Developed REST APIs using Django REST Framework for a healthcare patient management platform serving 200,000+ active users.",
                    "Optimized complex PostgreSQL queries using Indexing, Select Related, and Prefetch Related, reducing API response times by 40%.",
                    "Integrated third-party payment gateways (Stripe, Razorpay) with webhook validation and failure retry handlers."
                ]
            }
        ],
        "projects": [
            {
                "name": "Real-Time Bidding Platform Engine",
                "tech": "FastAPI, Redis Pub/Sub, PostgreSQL, Docker",
                "bullets": [
                    "Engineered an automated auction bidding system handling concurrent bids using Redis distributed locks and FastAPI WebSockets.",
                    "Created a dashboard monitoring system calculating top bids and user activity logs in real-time."
                ]
            },
            {
                "name": "Multi-Tenant SaaS Subscription Manager",
                "tech": "Django, Celery, PostgreSQL, AWS S3",
                "bullets": [
                    "Built a multi-tenant backend architecture handling automated subscription renewals, usage tracking, and billing invoice generation."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | NIT Trichy, India | Graduated: May 2020 | CGPA: 8.9/10.0"
        ],
        "certifications": [
            "AWS Certified Solutions Architect – Associate"
        ],
        "achievements": [
            "Speaker at PyCon India 2023 on 'High Performance Microservices with FastAPI and Asyncio'."
        ]
    },
    # 05 - Python Backend
    {
        "num": "05",
        "name": "Rohan Gupta",
        "role": "Python Backend Engineer",
        "summary": "Python Backend Developer with 2 years of professional experience building cloud backend services, microservices, and database models using Flask, FastAPI, Python, and MongoDB. Demonstrated expertise in containerization, API integration, and automated testing in fast-paced software environments.",
        "skills": [
            ("Languages", "Python, JavaScript (Basic), SQL"),
            ("Frameworks", "Flask, FastAPI, SQLAlchemy, Pytest"),
            ("Databases", "MongoDB, PostgreSQL, SQLite"),
            ("Tools & DevOps", "Docker, Docker Compose, Git, Postman, Linux, Nginx"),
            ("Web Technologies", "RESTful APIs, JSON, WebSockets, OAuth2")
        ],
        "experience": [
            {
                "title": "Python Backend Engineer",
                "company": "Innovex Digital, Gurgaon, India",
                "period": "Jun 2022 – Present",
                "bullets": [
                    "Built 20+ REST API endpoints using Flask and Flask-RESTful for an enterprise resource planning (ERP) system.",
                    "Designed schema-less data structures in MongoDB, facilitating dynamic field additions for inventory management.",
                    "Containerized development and production environments using Docker and Docker Compose, reducing environment configuration time for onboarding developers by 70%.",
                    "Wrote automated unit and integration tests using Pytest, achieving 84% test coverage across core application logic."
                ]
            },
            {
                "title": "Software Development Intern",
                "company": "DataCraft Labs, Delhi, India",
                "period": "Jan 2022 – May 2022",
                "bullets": [
                    "Assisted in building Python data ingestion scripts fetching JSON payloads from external REST APIs and loading them into PostgreSQL.",
                    "Implemented JWT authentication and password hashing (bcrypt) for user management modules."
                ]
            }
        ],
        "projects": [
            {
                "name": "IoT Device Fleet Monitoring Backend",
                "tech": "FastAPI, MongoDB, WebSockets, Docker",
                "bullets": [
                    "Developed a high-performance backend consuming telemetry data from 1,000+ simulated IoT sensors via WebSockets.",
                    "Configured automated alert notifications triggered when sensor metrics exceed safe thresholds."
                ]
            },
            {
                "name": "Content Management API with Role-Based Access Control",
                "tech": "Flask, SQLAlchemy, PostgreSQL",
                "bullets": [
                    "Created a modular Flask REST API managing articles, categories, user roles, and granular permission scopes."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Engineering | Delhi Technological University (DTU), Delhi, India | Graduated: May 2022 | Percentage: 79%"
        ],
        "certifications": [
            "Python Institute Certified Associate Programmer (PCAP)"
        ],
        "achievements": [
            "Contributed 5+ bug fixes to open-source Python Flask plugins."
        ]
    },
    # 06 - Python Backend
    {
        "num": "06",
        "name": "Sneha Kulkarni",
        "role": "Junior Python Developer",
        "summary": "Recent Computer Science graduate and Python Backend Intern with practical project experience building REST APIs, relational database schemas, and Python scripts. Solid understanding of Django, Python 3, object-oriented software design, and Git version control workflows.",
        "skills": [
            ("Languages", "Python, C++, HTML5, CSS3"),
            ("Web Frameworks", "Django, Django REST Framework, Flask"),
            ("Databases", "PostgreSQL, SQLite"),
            ("Tools", "Git, GitHub, Postman, VS Code, Linux Command Line"),
            ("Core Concepts", "Object-Oriented Programming (OOP), Data Structures, Algorithms, RESTful Web Services")
        ],
        "experience": [
            {
                "title": "Python Backend Intern",
                "company": "ByteCraft Technologies, Pune, India",
                "period": "Jan 2024 – Jun 2024",
                "bullets": [
                    "Developed RESTful APIs for an internal employee learning portal using Django REST Framework.",
                    "Created custom Django admin interfaces and database models, reducing administrative workflow effort by 30%.",
                    "Participated in weekly sprint reviews, pull request code reviews, and API testing using Postman."
                ]
            }
        ],
        "projects": [
            {
                "name": "Online Bookstore API",
                "tech": "Django, Django REST Framework, SQLite",
                "bullets": [
                    "Built a REST API featuring book search, filter by category, shopping cart management, and user order history.",
                    "Implemented pagination, filtering, and custom search backends for optimized catalog navigation."
                ]
            },
            {
                "name": "Task Automation CLI Utility",
                "tech": "Python, Click, SQLite",
                "bullets": [
                    "Developed a command-line interface tool allowing users to schedule, track, and categorize daily tasks with data persistence in SQLite."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science and Engineering | College of Engineering Pune (COEP), Pune, India | Graduated: Jun 2024 | CGPA: 8.4/10.0"
        ],
        "certifications": [
            "Meta Back-End Developer Specialization (Coursera)"
        ],
        "achievements": [
            "Secured 2nd rank in COEP Annual Coding Competition 2023."
        ]
    },

    # 07 - MERN Stack Developer
    {
        "num": "07",
        "name": "Vikramaditya Verma",
        "role": "Lead MERN Stack Developer",
        "summary": "Senior MERN Stack Developer with 5 years of full-stack development experience architecting high-scale web applications using React, Node.js, Express.js, TypeScript, and MongoDB. Proven success in leading frontend and backend software teams, optimizing web performance, and deploying cloud applications on AWS and Vercel.",
        "skills": [
            ("Frontend", "React 18, Next.js, Redux Toolkit, TypeScript, HTML5, CSS3, Tailwind CSS, Webpack"),
            ("Backend", "Node.js, Express.js, NestJS, REST APIs, GraphQL, Socket.io"),
            ("Databases", "MongoDB, Redis, Mongoose ORM, PostgreSQL"),
            ("DevOps & Cloud", "AWS (EC2, S3, CloudFront), Docker, CI/CD (GitHub Actions), Nginx"),
            ("Architecture", "Micro-frontends, Microservices, Monorepos (Turborepo), Serverless")
        ],
        "experience": [
            {
                "title": "Lead MERN Developer",
                "company": "UrbanTech Innovations, Bangalore, India",
                "period": "Feb 2022 – Present",
                "bullets": [
                    "Led the complete architectural redesign of a web-based e-learning platform using React 18, TypeScript, and Node.js microservices, scaling to 500,000+ active monthly users.",
                    "Built a real-time collaborative code editor and virtual classroom streaming interface utilizing Socket.io and WebRTC.",
                    "Improved Core Web Vitals across application pages, boosting Lighthouse performance scores from 62 to 96 through code splitting, dynamic imports, and image optimization.",
                    "Mentored a team of 6 engineers, managing sprint velocity, code quality standard enforcement, and architectural documentation."
                ]
            },
            {
                "title": "Full Stack Developer",
                "company": "Mindfire Solutions, Bhubaneswar, India",
                "period": "Jul 2019 – Jan 2022",
                "bullets": [
                    "Developed interactive web dashboards and REST APIs using React, Redux, Express.js, and MongoDB for e-commerce clients.",
                    "Built automated payment processing workflows integrating Stripe and PayPal APIs with transaction idempotency checks.",
                    "Implemented MongoDB aggregation pipelines to process analytical dashboards, accelerating report render speed by 50%."
                ]
            }
        ],
        "projects": [
            {
                "name": "Real-Time SaaS Project Collaboration Workspace",
                "tech": "React, Node.js, Socket.io, MongoDB, AWS S3",
                "bullets": [
                    "Built a Trello-like project management application featuring drag-and-drop task boards, live user presence indicators, and file attachments.",
                    "Configured automated JWT auth, refresh token rotation, and dynamic role permissions."
                ]
            },
            {
                "name": "Enterprise B2B E-Commerce Platform",
                "tech": "Next.js, TypeScript, Express.js, Redis, MongoDB",
                "bullets": [
                    "Architected a SSR/SSG e-commerce web application with full-text product search, dynamic multi-currency pricing, and Redis cart caching."
                ]
            }
        ],
        "education": [
            "B.Tech in Information Technology | KIIT University, Bhubaneswar, India | Graduated: May 2019 | CGPA: 8.8/10.0"
        ],
        "certifications": [
            "AWS Certified Developer – Associate",
            "MongoDB Certified Developer Associate"
        ],
        "achievements": [
            "Authored npm package 'react-canvas-draw-grid' with 15,000+ monthly downloads."
        ]
    },
    # 08 - MERN Stack Developer
    {
        "num": "08",
        "name": "Ananya Roy",
        "role": "Full Stack MERN Developer",
        "summary": "Full Stack MERN Developer with 3 years of experience building modern web applications with React, Node.js, Express, MongoDB, and TypeScript. Skilled in crafting responsive UI components, building secure RESTful APIs, and implementing state management architectures.",
        "skills": [
            ("Frontend", "React.js, JavaScript (ES6+), TypeScript, Redux Toolkit, HTML5, CSS3, Tailwind CSS, Bootstrap"),
            ("Backend", "Node.js, Express.js, RESTful APIs, JWT, Middleware"),
            ("Databases", "MongoDB, Mongoose, PostgreSQL"),
            ("Tools & Deployment", "Git, GitHub, Postman, Docker, Vercel, Render, Jest, RTL")
        ],
        "experience": [
            {
                "title": "MERN Stack Developer",
                "company": "PixelCraft Solutions, Kolkata, India",
                "period": "May 2022 – Present",
                "bullets": [
                    "Built 10+ web features for a SaaS HR management portal using React, TypeScript, Node.js, and MongoDB.",
                    "Created reusable component libraries using React and Tailwind CSS, standardizing UI design patterns across 4 internal projects.",
                    "Designed Express REST APIs with input validation (Joi/Zod) and error handling middleware, lowering API failure rates by 25%.",
                    "Integrated state management using Redux Toolkit and RTK Query for efficient server data fetching and caching."
                ]
            },
            {
                "title": "Junior Web Developer",
                "company": "WebTech India, Kolkata, India",
                "period": "Jun 2021 – Apr 2022",
                "bullets": [
                    "Developed responsive web landing pages and custom admin panels using React.js, HTML5, and CSS3.",
                    "Integrated backend REST APIs with Axios and implemented client-side form validation."
                ]
            }
        ],
        "projects": [
            {
                "name": "Healthcare Teleconsultation Portal",
                "tech": "React, Node.js, Express, MongoDB, WebRTC",
                "bullets": [
                    "Built a full-stack teleconsultation app enabling patients to schedule video appointments with doctors and receive digital prescriptions.",
                    "Integrated WebRTC for peer-to-peer secure video calls and Node.js backend for appointment management."
                ]
            },
            {
                "name": "Social Media Content Scheduling Application",
                "tech": "React, Express, MongoDB, Node-Cron",
                "bullets": [
                    "Developed an app allowing content creators to compose, preview, and automatically publish posts across social channels at scheduled intervals."
                ]
            }
        ],
        "education": [
            "B.Sc in Computer Science | St. Xavier's College, Kolkata, India | Graduated: May 2021 | Marks: 85%"
        ],
        "certifications": [
            "Meta Front-End Developer Professional Certificate"
        ],
        "achievements": [
            "Awarded 'Developer of the Quarter' at PixelCraft Solutions in Q3 2023."
        ]
    },
    # 09 - MERN Stack Developer
    {
        "num": "09",
        "name": "Karan Patel",
        "role": "Entry-Level MERN Developer",
        "summary": "Entry-Level MERN Stack Developer with 1 year of experience building responsive full-stack applications using React, Node.js, Express, and MongoDB. Strong knowledge of modern JavaScript (ES6+), HTML5/CSS3 layout design, REST API consumption, and Git version control.",
        "skills": [
            ("Frontend", "React.js, JavaScript, HTML5, CSS3, Bootstrap, Flexbox/Grid"),
            ("Backend", "Node.js, Express.js, REST APIs, JSON"),
            ("Databases", "MongoDB, Mongoose"),
            ("Tools", "Git, GitHub, VS Code, Postman, npm")
        ],
        "experience": [
            {
                "title": "Junior Full Stack Developer",
                "company": "AppDev Studio, Ahmedabad, India",
                "period": "Jul 2023 – Present",
                "bullets": [
                    "Developed frontend components in React for a customer feedback analytics tool.",
                    "Implemented backend API endpoints in Node.js and Express to fetch, parse, and write feedback records into MongoDB.",
                    "Fixed 30+ UI responsiveness and cross-browser compatibility bugs reported across Chrome, Firefox, and Safari.",
                    "Participated in daily standups, code reviews, and feature deployment pipelines."
                ]
            }
        ],
        "projects": [
            {
                "name": "Event Management & Ticket Booking Web App",
                "tech": "React, Node.js, Express, MongoDB",
                "bullets": [
                    "Created a web application where users can discover local events, select ticket tiers, and view virtual QR tickets.",
                    "Implemented user registration, password hashing (bcrypt), and JWT session persistence."
                ]
            },
            {
                "name": "Personal Portfolio & Tech Blog System",
                "tech": "React, Markdown, Express, MongoDB",
                "bullets": [
                    "Developed a personal portfolio application featuring a dynamic blog reader consuming Markdown content via Express API."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Engineering | Gujarat Technological University, Ahmedabad, India | Graduated: May 2023 | CGPA: 7.9/10.0"
        ],
        "certifications": [
            "FreeCodeCamp Full Stack Web Development Certification"
        ],
        "achievements": [
            "Built open-source React component library used in university project portal."
        ]
    },

    # 10 - Full Stack Developer (.NET)
    {
        "num": "10",
        "name": "Siddharth Nambiar",
        "role": "Senior .NET Full Stack Engineer",
        "summary": "Senior .NET Full Stack Architect with 5+ years of experience engineering enterprise web applications and cloud services using C#, ASP.NET Core, Angular, SQL Server, and Azure. Proven expertise in microservices design, Entity Framework Core optimization, and CI/CD automation.",
        "skills": [
            ("Languages", "C#, SQL, TypeScript, JavaScript"),
            ("Backend Frameworks", ".NET 8 / ASP.NET Core, Web API, Entity Framework Core, LINQ, SignalR"),
            ("Frontend Frameworks", "Angular 16/14, RxJS, NgRx, HTML5, SCSS, Bootstrap"),
            ("Databases", "Microsoft SQL Server, Azure SQL, Redis"),
            ("Cloud & DevOps", "Azure (App Services, Azure Functions, Blob Storage), Docker, Azure DevOps, CI/CD Pipelines")
        ],
        "experience": [
            {
                "title": "Senior .NET Engineer",
                "company": "Cognizant Technology Solutions, Chennai, India",
                "period": "Jan 2021 – Present",
                "bullets": [
                    "Architected enterprise healthcare solutions using ASP.NET Core Web API, EF Core, Angular 16, and SQL Server serving 1M+ active patients.",
                    "Migrated monolithic .NET Framework applications to cloud-native .NET 7 microservices on Azure App Services, improving system reliability to 99.95%.",
                    "Optimized SQL Server stored procedures, indexing strategies, and EF Core LINQ queries, reducing core report generation time by 60%.",
                    "Configured CI/CD pipelines in Azure DevOps for automated building, testing, and multi-stage production deployment."
                ]
            },
            {
                "title": "Software Developer",
                "company": "TCS (Tata Consultancy Services), Chennai, India",
                "period": "Jun 2018 – Dec 2020",
                "bullets": [
                    "Developed web modules for a financial insurance platform using C#, ASP.NET MVC, and SQL Server.",
                    "Implemented real-time claim status notification engines using SignalR and Angular UI components."
                ]
            }
        ],
        "projects": [
            {
                "name": "Enterprise Claims Processing Portal",
                "tech": ".NET 8, ASP.NET Core, Angular, SQL Server, Azure",
                "bullets": [
                    "Designed an automated insurance claims engine processing 100,000+ monthly policy transactions.",
                    "Integrated Azure Blob Storage for secure medical document uploads with granular SAS token security."
                ]
            },
            {
                "name": "Real-Time Telemetry & Monitoring Dashboard",
                "tech": "ASP.NET Core SignalR, C#, Angular, Redis",
                "bullets": [
                    "Built a live telemetry stream application pushing server health metrics to Angular dashboards via WebSockets."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | Anna University, Chennai, India | Graduated: May 2018 | CGPA: 8.6/10.0"
        ],
        "certifications": [
            "Microsoft Certified: Azure Developer Associate (AZ-204)",
            "Microsoft Certified: DevOps Engineer Expert (AZ-400)"
        ],
        "achievements": [
            "Recognized as 'Best Architect of the Year' at Cognizant in 2023 for cloud modernization work."
        ]
    },
    # 11 - Full Stack Developer (.NET)
    {
        "num": "11",
        "name": "Manish Sharma",
        "role": ".NET Software Developer",
        "summary": "Full Stack .NET Developer with 3 years of software engineering experience developing web applications using C#, ASP.NET Core, React, and SQL Server. Skilled in REST API development, Object-Relational Mapping (EF Core), and modern frontend component design.",
        "skills": [
            ("Languages", "C#, JavaScript, TypeScript, SQL"),
            ("Backend", "ASP.NET Core Web API, Entity Framework Core, LINQ, REST APIs"),
            ("Frontend", "React.js, Redux, HTML5, CSS3, Tailwind CSS"),
            ("Databases", "Microsoft SQL Server, PostgreSQL"),
            ("Tools", "Visual Studio 2022, Git, Postman, Azure DevOps, Swagger")
        ],
        "experience": [
            {
                "title": "Software Developer (.NET)",
                "company": "LTIMindtree, Mumbai, India",
                "period": "Aug 2021 – Present",
                "bullets": [
                    "Developed secure RESTful Web APIs using ASP.NET Core 6 and C# for an inventory management system.",
                    "Built responsive frontend interfaces in React with Redux state management, connecting seamlessly to .NET backend APIs.",
                    "Implemented identity authentication using ASP.NET Core Identity and JWT tokens with refresh token handlers.",
                    "Wrote unit tests using xUnit and Moq, achieving 82% code coverage across business logic layers."
                ]
            }
        ],
        "projects": [
            {
                "name": "B2B Vendor Management System",
                "tech": "ASP.NET Core, C#, React, EF Core, SQL Server",
                "bullets": [
                    "Engineered a vendor portal for invoice submission, purchase order tracking, and vendor evaluation scoring.",
                    "Implemented automated background batch processing using Quartz.NET for overnight report generation."
                ]
            },
            {
                "name": "Asset Tracking Web Application",
                "tech": "ASP.NET Core Web API, React, SQL Server",
                "bullets": [
                    "Created an internal IT asset allocation system with barcode scanner integration and audit trail reporting."
                ]
            }
        ],
        "education": [
            "B.E. in Information Technology | Mumbai University, Mumbai, India | Graduated: Jun 2021 | CGPA: 8.2/10.0"
        ],
        "certifications": [
            "Developing Solutions for Microsoft Azure (AZ-204) Trained"
        ],
        "achievements": [
            "Received Excellence Award for delivering vendor portal ahead of client timeline."
        ]
    },
    # 12 - Full Stack Developer (.NET)
    {
        "num": "12",
        "name": "Abhishek Joshi",
        "role": "Junior .NET Developer",
        "summary": "Junior .NET Developer with 1.5 years of experience developing enterprise web solutions with C#, ASP.NET Core, Entity Framework Core, SQL Server, and JavaScript. Solid understanding of Object-Oriented Programming, MVC pattern, and database design.",
        "skills": [
            ("Languages", "C#, SQL, JavaScript, HTML5, CSS3"),
            ("Frameworks", "ASP.NET Core MVC, ASP.NET Web API, Entity Framework Core"),
            ("Databases", "MS SQL Server, SSMS"),
            ("Tools", "Visual Studio, Git, GitHub, Azure DevOps, Postman")
        ],
        "experience": [
            {
                "title": "Junior Software Engineer",
                "company": "Hexaware Technologies, Chennai, India",
                "period": "Feb 2023 – Present",
                "bullets": [
                    "Maintained and updated web applications built on ASP.NET Core MVC and SQL Server for insurance clients.",
                    "Created dynamic web forms and views using Razor Syntax, HTML5, CSS3, and JavaScript.",
                    "Wrote T-SQL queries, views, and stored procedures to handle business data processing.",
                    "Participated in bug resolution sprints, resolving 40+ tickets related to UI display and database validation errors."
                ]
            }
        ],
        "projects": [
            {
                "name": "Employee Expense Management System",
                "tech": "ASP.NET Core MVC, EF Core, SQL Server, Bootstrap",
                "bullets": [
                    "Built a web app enabling employees to log travel expenses, upload receipt images, and track manager approvals.",
                    "Implemented role-based authorization for Employee, Manager, and Finance Admin accounts."
                ]
            },
            {
                "name": "Library Book Reservation Portal",
                "tech": "C#, ASP.NET Core Web API, SQL Server",
                "bullets": [
                    "Developed REST endpoints for searching catalog books, issuing digital library cards, and checking due dates."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | SRM Institute of Science and Technology, Chennai, India | Graduated: Dec 2022 | CGPA: 8.0/10.0"
        ],
        "certifications": [
            "Foundational C# Certification with FreeCodeCamp"
        ],
        "achievements": [
            "Graduated with First Class Distinction from SRM IST."
        ]
    },

    # 13 - React Frontend Developer
    {
        "num": "13",
        "name": "Divya Nair",
        "role": "Senior React Frontend Developer",
        "summary": "Senior Frontend Developer with 4.5 years of experience building scalable, performant, and accessible web applications using React 18, Next.js, TypeScript, Redux Toolkit, and modern CSS architectures. Proven ability to optimize application performance, establish UI component design systems, and mentor engineering teams.",
        "skills": [
            ("Frontend Core", "React 18, Next.js 14, TypeScript, JavaScript (ES6+), HTML5, CSS3, SASS"),
            ("State & Data Fetching", "Redux Toolkit, RTK Query, React Query (TanStack), Context API"),
            ("Styling & UI", "Tailwind CSS, Material UI, Styled Components, Framer Motion"),
            ("Testing & Performance", "Jest, React Testing Library, Cypress, Lighthouse, Core Web Vitals"),
            ("Tools & Build", "Webpack, Vite, npm/pnpm, Git, GitHub Actions, Figma")
        ],
        "experience": [
            {
                "title": "Senior Frontend Engineer",
                "company": "Paytm, Bangalore, India",
                "period": "Oct 2022 – Present",
                "bullets": [
                    "Architected merchant-facing dashboard using React 18, TypeScript, and Tailwind CSS, serving 2M+ active daily business owners.",
                    "Improved page load speed by 55% and reduced bundle size by 180KB using code splitting, dynamic imports, and lazy loading techniques.",
                    "Built a reusable UI Design System component library used across 8 internal web projects, ensuring WCAG 2.1 AA accessibility compliance.",
                    "Established unit testing standards using Jest and React Testing Library, pushing test coverage from 45% to 86%."
                ]
            },
            {
                "title": "Frontend Developer",
                "company": "MakeMyTrip, Gurgaon, India",
                "period": "Jul 2019 – Sep 2022",
                "bullets": [
                    "Developed high-conversion flight and hotel booking web pages using Next.js, React, and Redux.",
                    "Implemented server-side rendering (SSR) and dynamic metadata generation, increasing organic search SEO traffic by 35%."
                ]
            }
        ],
        "projects": [
            {
                "name": "Cryptocurrency Analytics & Live Trading Dashboard",
                "tech": "React 18, TypeScript, Tailwind CSS, Recharts, WebSockets",
                "bullets": [
                    "Built a real-time crypto price tracking dashboard with live candlestick charts consuming WebSocket ticker feeds.",
                    "Implemented customizable widget dashboards with localized state storage."
                ]
            },
            {
                "name": "SaaS Design System & Component Studio",
                "tech": "React, Storybook, TypeScript, Tailwind CSS",
                "bullets": [
                    "Created an interactive component story library documenting 50+ accessible frontend UI controls."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | Kerala University, Trivandrum, India | Graduated: May 2019 | CGPA: 8.6/10.0"
        ],
        "certifications": [
            "Meta Frontend Developer Professional Certificate"
        ],
        "achievements": [
            "Featured speaker at React India Conference 2023."
        ]
    },
    # 14 - React Frontend Developer
    {
        "num": "14",
        "name": "Meera Krishnan",
        "role": "React Frontend Engineer",
        "summary": "Frontend Web Developer with 2.5 years of experience crafting interactive, user-centered web applications using React, JavaScript (ES6+), HTML5, CSS3, Tailwind CSS, and REST APIs. Passionate about responsive UI design, component reusable architecture, and web performance optimization.",
        "skills": [
            ("Core Frontend", "React.js, JavaScript (ES6+), HTML5, CSS3, Tailwind CSS, Bootstrap"),
            ("State Management", "React Context API, Redux Toolkit, Zustand"),
            ("APIs & Tools", "REST APIs, Axios, Git, GitHub, VS Code, Vite, Postman"),
            ("Design & UX", "Responsive Web Design, Figma to Code, CSS Modules, Flexbox/Grid")
        ],
        "experience": [
            {
                "title": "Frontend Developer",
                "company": "InMobi, Bangalore, India",
                "period": "Jan 2022 – Present",
                "bullets": [
                    "Developed dynamic web user interfaces for an ad analytics campaign manager using React.js and Tailwind CSS.",
                    "Transformed complex Figma wireframes into pixel-perfect, mobile-responsive web pages across modern browsers.",
                    "Integrated REST APIs using Axios and Zustand state management, handling seamless data display and updates.",
                    "Reduced DOM rendering lag on complex analytics data tables by implementing React virtualized list rendering."
                ]
            },
            {
                "title": "Frontend Web Developer Intern",
                "company": "Zeta Suite, Bangalore, India",
                "period": "Jun 2021 – Dec 2021",
                "bullets": [
                    "Assisted in creating landing pages and promotional web UI widgets using React, HTML5, and SASS.",
                    "Resolved cross-browser layout inconsistencies and accessibility issues."
                ]
            }
        ],
        "projects": [
            {
                "name": "Streaming Video Platform Web UI",
                "tech": "React, Tailwind CSS, TMDB API, React Router",
                "bullets": [
                    "Built a Netflix-style video discovery application featuring movie browsing, trailer popups, search filters, and watchlist management."
                ]
            },
            {
                "name": "Kanban Task Management Board",
                "tech": "React, React-Beautiful-DND, Context API, CSS Modules",
                "bullets": [
                    "Developed an intuitive drag-and-drop task organizer application with customizable columns and local storage persistence."
                ]
            }
        ],
        "education": [
            "B.Sc in Information Technology | Calicut University, Kerala, India | Graduated: May 2021 | Marks: 81%"
        ],
        "certifications": [
            "Frontend Web Development Specialist (Udacity)"
        ],
        "achievements": [
            "Winner of Best UI Design award at InMobi Internal Hackathon 2023."
        ]
    },
    # 15 - React Frontend Developer
    {
        "num": "15",
        "name": "Rahul Verma",
        "role": "Junior React Developer",
        "summary": "Enthusiastic Junior React Developer with 1 year of web development experience building single-page applications (SPAs) with React, HTML5, CSS3, JavaScript, and Bootstrap. Dedicated to creating clean code, intuitive UI layouts, and accessible frontend web experiences.",
        "skills": [
            ("Languages", "JavaScript (ES6+), HTML5, CSS3, C++"),
            ("Frontend", "React.js, Bootstrap, Flexbox, Grid, CSS Modules"),
            ("Tools", "Git, GitHub, VS Code, npm, Postman, Vite"),
            ("Concepts", "Component Lifecycle, React Hooks, State & Props, DOM Manipulation")
        ],
        "experience": [
            {
                "title": "Junior Frontend Web Developer",
                "company": "WebSpectrum Solutions, Noida, India",
                "period": "Jul 2023 – Present",
                "bullets": [
                    "Developed 15+ custom web components in React for client corporate websites.",
                    "Collaborated with backend engineers to integrate RESTful JSON endpoints into frontend React states.",
                    "Tested and optimized web applications for mobile, tablet, and desktop viewport responsiveness.",
                    "Managed version control pull requests and code merges via Git and GitHub."
                ]
            }
        ],
        "projects": [
            {
                "name": "Weather & Air Quality Web App",
                "tech": "React, OpenWeatherMap API, CSS3, Vite",
                "bullets": [
                    "Built a responsive weather web app displaying 7-day weather forecasts, air quality index, and geolocation search."
                ]
            },
            {
                "name": "E-Commerce Product Catalog UI",
                "tech": "React, Context API, Bootstrap 5",
                "bullets": [
                    "Created a shopping application UI featuring product filtering by price/category, dynamic cart badge counter, and checkout modal."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | Amity University, Noida, India | Graduated: Jun 2023 | CGPA: 7.8/10.0"
        ],
        "certifications": [
            "React Basics Certification (Coursera - Meta)"
        ],
        "achievements": [
            "Represented college in National Smart India Hackathon 2022."
        ]
    },

    # 16 - Machine Learning Engineer
    {
        "num": "16",
        "name": "Siddharth Rao",
        "role": "Senior Machine Learning Engineer",
        "summary": "Senior Machine Learning Engineer with 4.5 years of experience building, deploying, and monitoring end-to-end ML models in production environments. Proficient in Python, PyTorch, TensorFlow, Scikit-Learn, MLOps, MLflow, and cloud deployments on AWS. Specialized in Computer Vision, Predictive Analytics, and Model Optimization.",
        "skills": [
            ("ML/DL Frameworks", "PyTorch, TensorFlow, Keras, Scikit-Learn, XGBoost, OpenCV"),
            ("Languages", "Python, C++, SQL, R"),
            ("MLOps & Data Pipeline", "MLflow, Kubeflow, Ray, DVC, Airflow, Docker, AWS SageMaker"),
            ("Math & Theory", "Deep Learning, Linear Algebra, Probability & Statistics, Optimization"),
            ("Databases", "PostgreSQL, MongoDB, Redis, Snowflake")
        ],
        "experience": [
            {
                "title": "Senior ML Engineer",
                "company": "Fractal Analytics, Mumbai, India",
                "period": "Oct 2022 – Present",
                "bullets": [
                    "Designed and deployed real-time fraud detection models using XGBoost and PyTorch, analyzing 5M+ daily transaction logs with 98.4% precision.",
                    "Established MLOps infrastructure using MLflow, DVC, and AWS SageMaker, automating model retraining pipelines and reducing deployment cycles from 2 weeks to 2 hours.",
                    "Optimized deep learning model inference speed by 3.5x using TensorRT quantization and ONNX runtime optimization for edge deployment.",
                    "Led a squad of 3 data scientists, overseeing dataset curation, feature engineering pipelines, and model evaluation metrics."
                ]
            },
            {
                "title": "Machine Learning Engineer",
                "company": "Tiger Analytics, Chennai, India",
                "period": "Jul 2019 – Sep 2022",
                "bullets": [
                    "Built customer churn prediction models using Scikit-Learn and LightGBM for telecommunication enterprises, improving retention rates by 18%.",
                    "Engineered automated computer vision defect detection systems for industrial manufacturing lines using OpenCV and PyTorch."
                ]
            }
        ],
        "projects": [
            {
                "name": "Automated Medical Image Segmentation System",
                "tech": "PyTorch, U-Net, OpenCV, FastAPI, Docker",
                "bullets": [
                    "Developed a deep learning model for automated MRI lesion segmentation achieving 0.91 Dice Similarity Coefficient.",
                    "Deployed model endpoint as a containerized FastAPI web service integrated with hospital PACS servers."
                ]
            },
            {
                "name": "Real-Time Demand Forecasting Engine",
                "tech": "Python, Prophet, XGBoost, MLflow, AWS SageMaker",
                "bullets": [
                    "Built a time-series forecasting model predicting inventory demand across 500+ retail stores with 94% accuracy."
                ]
            }
        ],
        "education": [
            "M.Tech in Artificial Intelligence | IIT Madras, Chennai, India | Graduated: May 2019 | CGPA: 9.1/10.0",
            "B.Tech in Computer Science | Anna University, Chennai, India | Graduated: May 2017 | CGPA: 8.7/10.0"
        ],
        "certifications": [
            "AWS Certified Machine Learning – Specialty",
            "TensorFlow Developer Certificate"
        ],
        "achievements": [
            "Kaggle Competition Master (Top 1% ranking in 2 international ML challenges)."
        ]
    },
    # 17 - Machine Learning Engineer
    {
        "num": "17",
        "name": "Nikhil Agarwal",
        "role": "Machine Learning Engineer",
        "summary": "Machine Learning Engineer with 2.5 years of experience building statistical models, predictive analytics algorithms, and machine learning pipelines using Python, Scikit-learn, XGBoost, and FastAPI. Passionate about feature engineering, data validation, and building production-ready ML REST services.",
        "skills": [
            ("Languages", "Python, SQL, R"),
            ("Libraries", "Scikit-learn, XGBoost, Pandas, NumPy, SciPy, Matplotlib, Seaborn"),
            ("ML Frameworks", "PyTorch (Basic), LightGBM, Statsmodels"),
            ("Deployment & Tools", "FastAPI, Flask, Docker, Git, MLflow, Jupyter, Postman")
        ],
        "experience": [
            {
                "title": "Machine Learning Engineer",
                "company": "Mu Sigma, Bangalore, India",
                "period": "Jan 2022 – Present",
                "bullets": [
                    "Developed predictive lead scoring machine learning models using Scikit-Learn and Random Forest algorithms, increasing sales conversion rate by 22%.",
                    "Built feature engineering pipelines processing 100GB+ tabular datasets using Pandas and SQL queries.",
                    "Deployed trained ML models as REST APIs using FastAPI and Docker containers deployed on AWS EC2.",
                    "Monitored production model performance for feature drift and accuracy metrics using custom logging dashboards."
                ]
            },
            {
                "title": "Data Science Intern",
                "company": "Analytics Vidhya, Gurgaon, India",
                "period": "Jun 2021 – Dec 2021",
                "bullets": [
                    "Cleaned, preprocessed, and analyzed multi-source consumer data for EDA and baseline model benchmarks.",
                    "Wrote automated data validation scripts checking missing value rates and anomaly scores."
                ]
            }
        ],
        "projects": [
            {
                "name": "Credit Risk Assessment & Default Prediction Model",
                "tech": "Python, XGBoost, Scikit-learn, FastAPI, Docker",
                "bullets": [
                    "Engineered a loan default risk classifier utilizing SMOTE oversampling for imbalanced credit datasets.",
                    "Achieved 0.88 ROC-AUC score and built FastAPI interface for instant loan applicant risk scoring."
                ]
            },
            {
                "name": "E-Commerce Customer Recommendation Engine",
                "tech": "Python, Collaborative Filtering, Implicit ALS, Flask",
                "bullets": [
                    "Developed a personalized item recommendation algorithm for e-commerce shoppers based on implicit clickstream behaviors."
                ]
            }
        ],
        "education": [
            "B.Tech in Applied Mathematics & Computing | Delhi Technological University (DTU), India | Graduated: May 2021 | CGPA: 8.3/10.0"
        ],
        "certifications": [
            "Deep Learning Specialization by Andrew Ng (Coursera)"
        ],
        "achievements": [
            "Secured 3rd Rank in National Data Science Hackathon 2022."
        ]
    },
    # 18 - Machine Learning Engineer
    {
        "num": "18",
        "name": "Pooja Hegde",
        "role": "Junior Machine Learning Engineer",
        "summary": "Recent Computer Science graduate with hands-on project experience in Machine Learning, Python programming, Data Cleaning, and Computer Vision. Solid understanding of regression, classification algorithms, neural networks, and OpenCV.",
        "skills": [
            ("Languages", "Python, C++, SQL"),
            ("ML Tools", "Scikit-learn, OpenCV, TensorFlow (Keras), Pandas, NumPy, Matplotlib"),
            ("Concepts", "Supervised & Unsupervised Learning, Feature Scaling, Cross-Validation, CNNs"),
            ("Tools", "Jupyter Notebooks, VS Code, Git, GitHub, Google Colab")
        ],
        "experience": [
            {
                "title": "Machine Learning Intern",
                "company": "SmartData Solutions, Hyderabad, India",
                "period": "Jan 2024 – Jun 2024",
                "bullets": [
                    "Preprocessed tabular datasets and engineered domain features for a real estate price prediction project.",
                    "Trained and evaluated baseline models using Linear Regression, Decision Trees, and Random Forest in Scikit-learn.",
                    "Visualized model performance metrics, confusion matrices, and feature importance graphs for stakeholder review."
                ]
            }
        ],
        "projects": [
            {
                "name": "Automated License Plate Recognition (ALPR)",
                "tech": "Python, OpenCV, EasyOCR, Flask",
                "bullets": [
                    "Built an image processing pipeline detecting vehicle license plates from video feeds and extracting text using OCR.",
                    "Achieved 91% plate text recognition accuracy across test video footage."
                ]
            },
            {
                "name": "Customer Sentiment Analysis on Product Reviews",
                "tech": "Python, NLTK, TF-IDF, Logistic Regression, Streamlit",
                "bullets": [
                    "Trained an NLP sentiment classifier categorizing user text reviews into Positive, Negative, or Neutral with 86% accuracy."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science and Engineering | BITS Pilani (Hyderabad Campus), India | Graduated: Jun 2024 | CGPA: 8.2/10.0"
        ],
        "certifications": [
            "Machine Learning Specialization by Stanford Online (Coursera)"
        ],
        "achievements": [
            "Presented capstone ALPR project at BITS Annual Tech Symposium."
        ]
    },

    # 19 - AI / Generative AI Engineer
    {
        "num": "19",
        "name": "Aditya Sengupta",
        "role": "Senior Generative AI Engineer",
        "summary": "Senior AI Architect with 4.5+ years of software experience specializing in Generative AI, Large Language Models (LLMs), Retrieval-Augmented Generation (RAG), and Agentic Workflows. Expertise in building production RAG platforms using LangChain, LangGraph, Vector Databases (Pinecone, ChromaDB), OpenAI API, and Llama 3.",
        "skills": [
            ("GenAI Core", "LLMs (GPT-4, Llama 3, Claude 3), LangChain, LangGraph, LlamaIndex, AutoGen"),
            ("RAG Architecture", "Vector DBs (Pinecone, ChromaDB, FAISS, Milvus), Semantic Search, Hybrid Search, Reranking"),
            ("Languages & Backend", "Python 3.11, FastAPI, Asyncio, SQL, TypeScript"),
            ("ML Frameworks", "PyTorch, Hugging Face Transformers, PEFT, LoRA, Sentence-Transformers"),
            ("Cloud & MLOps", "AWS (SageMaker, Bedrock), Docker, LangSmith, Weights & Biases, Git")
        ],
        "experience": [
            {
                "title": "Lead Generative AI Engineer",
                "company": "Kritikal AI Systems, Bangalore, India",
                "period": "Nov 2022 – Present",
                "bullets": [
                    "Architected enterprise RAG system processing 100,000+ internal PDF/Word documents using LangChain, OpenAI embeddings, and Pinecone, cutting document search time by 80%.",
                    "Built multi-agent conversational AI workflows using LangGraph and AutoGen for automated customer support ticket triage.",
                    "Fine-tuned Llama 3 8B parameter models using LoRA and PEFT on domain-specific legal datasets, improving legal entity extraction accuracy by 24%.",
                    "Implemented LangSmith tracing and evaluation metrics to monitor hallucination rates, latency, and token consumption."
                ]
            },
            {
                "title": "AI / ML Developer",
                "company": "Persistent Systems, Pune, India",
                "period": "Jul 2019 – Oct 2022",
                "bullets": [
                    "Developed NLP text summarization and named entity recognition (NER) pipelines using Hugging Face Transformers and SpaCy.",
                    "Built FastAPI backends serving machine learning and text processing models to mobile web applications."
                ]
            }
        ],
        "projects": [
            {
                "name": "Enterprise Knowledge Graph RAG Copilot",
                "tech": "Python, LangGraph, Neo4j, Pinecone, GPT-4, FastAPI",
                "bullets": [
                    "Engineered a GraphRAG system combining structured knowledge graphs with vector similarity search for complex enterprise Q&A.",
                    "Built conversational memory, dynamic intent routing, and dynamic source citation attribution."
                ]
            },
            {
                "name": "Autonomous Code Review Agent",
                "tech": "Python, LangChain, Claude API, GitHub API, Docker",
                "bullets": [
                    "Developed an automated agent inspecting GitHub Pull Requests, highlighting code security vulnerabilities and suggesting refactoring diffs."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | IIIT Hyderabad, India | Graduated: May 2019 | CGPA: 8.9/10.0"
        ],
        "certifications": [
            "DeepLearning.AI Generative AI with LLMs Certificate",
            "AWS Certified Solutions Architect – Associate"
        ],
        "achievements": [
            "Created open-source LangChain plugin with 1,200+ GitHub Stars."
        ]
    },
    # 20 - AI / Generative AI Engineer
    {
        "num": "20",
        "name": "Kavya Menon",
        "role": "Generative AI Engineer",
        "summary": "Generative AI Engineer with 2 years of software development experience creating LLM-powered applications, RAG pipelines, and conversational chatbots. Proficient in Python, Hugging Face, LlamaIndex, ChromaDB, Prompt Engineering, and FastAPI.",
        "skills": [
            ("AI / GenAI", "OpenAI API, Anthropic API, LangChain, LlamaIndex, Hugging Face, Prompt Engineering"),
            ("Vector Databases", "ChromaDB, FAISS, Pinecone"),
            ("Languages & Web", "Python, FastAPI, Streamlit, SQL, JSON"),
            ("Tools", "Docker, Git, GitHub, Postman, Jupyter Notebooks")
        ],
        "experience": [
            {
                "title": "Generative AI Developer",
                "company": "Tredence Analytics, Bangalore, India",
                "period": "Jul 2022 – Present",
                "bullets": [
                    "Built RAG-based AI assistant for internal company HR policies using LlamaIndex, OpenAI embeddings, and ChromaDB vector store.",
                    "Implemented advanced prompt engineering techniques (Few-Shot, Chain-of-Thought) to improve structured JSON output consistency from LLMs to 96%.",
                    "Developed FastAPI microservices handling asynchronous streaming responses (SSE) to Streamlit frontend interfaces.",
                    "Optimized token cost consumption by implementing vector caching strategies with FAISS, lowering monthly API expenses by 35%."
                ]
            }
        ],
        "projects": [
            {
                "name": "Document Intelligence & Q&A Assistant",
                "tech": "Python, LangChain, OpenAI, ChromaDB, Streamlit",
                "bullets": [
                    "Created an interactive web app enabling users to upload PDF contracts and ask complex semantic questions with exact page references."
                ]
            },
            {
                "name": "AI Email Response Generator",
                "tech": "Python, Hugging Face, FastAPI, React",
                "bullets": [
                    "Built a browser extension backend generating contextual email reply suggestions based on thread sentiment and tone settings."
                ]
            }
        ],
        "education": [
            "B.Tech in Information Technology | Amrita Vishwa Vidyapeetham, Coimbatore, India | Graduated: May 2022 | CGPA: 8.3/10.0"
        ],
        "certifications": [
            "LangChain Developer Certification (ActiveLoop)"
        ],
        "achievements": [
            "Won 1st prize at Bangalore GenAI Hackathon 2023."
        ]
    },
    # 21 - AI / Generative AI Engineer
    {
        "num": "21",
        "name": "Varun Choudhary",
        "role": "AI Application Developer",
        "summary": "AI Software Developer with 1 year of experience building Python web services, vector search applications, and OpenAI API integrations. Good foundational knowledge of NLP, embeddings, RAG concepts, and web microservices.",
        "skills": [
            ("AI & NLP", "OpenAI API, LangChain, Sentence-Transformers, Hugging Face, Prompt Design"),
            ("Vector DB", "ChromaDB, FAISS"),
            ("Programming", "Python, JavaScript, SQL"),
            ("Frameworks & Tools", "FastAPI, Flask, Streamlit, Git, Docker")
        ],
        "experience": [
            {
                "title": "Junior AI Developer",
                "company": "Cognitive Cloud Labs, Hyderabad, India",
                "period": "Aug 2023 – Present",
                "bullets": [
                    "Integrated OpenAI and Anthropic API endpoints into existing enterprise Python backend platforms.",
                    "Built document chunking and vector indexing scripts using LangChain text splitters and ChromaDB.",
                    "Created Streamlit dashboard prototypes demonstrating AI chatbot capabilities to prospective clients.",
                    "Assisted in maintaining unit tests and API documentation for internal GenAI tools."
                ]
            }
        ],
        "projects": [
            {
                "name": "Code Base Search & Explanation Bot",
                "tech": "Python, LangChain, ChromaDB, OpenAI API",
                "bullets": [
                    "Built a tool indexing Git source code repositories and allowing developers to ask natural language questions about codebase architecture."
                ]
            },
            {
                "name": "Smart Customer Review Summarizer",
                "tech": "Python, Hugging Face Transformers, Flask",
                "bullets": [
                    "Developed a web API summarizing hundreds of user product feedback reviews into key pros, cons, and actionable bullet points."
                ]
            }
        ],
        "education": [
            "B.E. in Computer Science | Osmania University, Hyderabad, India | Graduated: Jun 2023 | CGPA: 7.9/10.0"
        ],
        "certifications": [
            "Building Systems with the ChatGPT API (DeepLearning.AI)"
        ],
        "achievements": [
            "Organized 24-hour AI Hackathon at Osmania University."
        ]
    },

    # 22 - Data Engineer
    {
        "num": "22",
        "name": "Harish Iyer",
        "role": "Senior Data Engineer",
        "summary": "Senior Data Engineer with 5+ years of experience designing, constructing, and managing enterprise data platforms, data warehouses, and real-time ETL pipelines. Expertise in PySpark, Apache Kafka, Apache Airflow, Snowflake, Databricks, and AWS big data services.",
        "skills": [
            ("Big Data & Processing", "Apache Spark / PySpark, Apache Kafka, Hadoop, MapReduce, Delta Lake"),
            ("Orchestration & ETL", "Apache Airflow, dbt, AWS Glue, NiFi, Luigi"),
            ("Data Warehousing", "Snowflake, Amazon Redshift, Google BigQuery, Databricks"),
            ("Languages", "Python, SQL, Scala, Bash"),
            ("Cloud & Infrastructure", "AWS (S3, EMR, Athena, Lambda), Docker, Terraform, Git")
        ],
        "experience": [
            {
                "title": "Senior Data Engineer",
                "company": "Swiggy, Bangalore, India",
                "period": "Jan 2022 – Present",
                "bullets": [
                    "Architected real-time streaming telemetry pipeline using Apache Kafka and PySpark, processing 200M+ daily events into Snowflake.",
                    "Optimized Airflow DAG schedules and Spark memory management configurations, cutting nightly ETL execution duration by 45%.",
                    "Implemented Medallion Architecture (Bronze/Silver/Gold layers) on Databricks Delta Lake, ensuring high data reliability for business intelligence teams.",
                    "Designed data governance and access control policies across Snowflake schemas using role-based security."
                ]
            },
            {
                "title": "Data Engineer",
                "company": "Mu Sigma, Bangalore, India",
                "period": "Jun 2018 – Dec 2021",
                "bullets": [
                    "Built batch processing pipelines using Python, SQL, and AWS Glue, loading structured data into Amazon Redshift data warehouses.",
                    "Wrote data quality monitoring tests with Great Expectations, preventing corrupt data loads into production dashboards."
                ]
            }
        ],
        "projects": [
            {
                "name": "Financial Transaction Streaming Analytics Pipeline",
                "tech": "PySpark, Kafka, Snowflake, Airflow, AWS S3",
                "bullets": [
                    "Engineered an end-to-end streaming data platform processing real-time credit card swipe data for fraud detection and daily reconciliation.",
                    "Constructed automated data validation and dead-letter queues (DLQ) for malformed incoming payloads."
                ]
            },
            {
                "name": "Enterprise Data Warehouse Modernization",
                "tech": "dbt, Snowflake, Airflow, GitHub Actions",
                "bullets": [
                    "Migrated 300+ legacy SQL stored procedures to modular dbt data transformations on Snowflake with automated CI/CD testing."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | RV College of Engineering, Bangalore, India | Graduated: May 2018 | CGPA: 8.8/10.0"
        ],
        "certifications": [
            "Databricks Certified Data Engineer Professional",
            "Snowflake SnowPro Core Certification"
        ],
        "achievements": [
            "Awarded Swiggy Tech Star Award 2023 for platform performance optimization."
        ]
    },
    # 23 - Data Engineer
    {
        "num": "23",
        "name": "Gautam Deshmukh",
        "role": "Data Pipeline Engineer",
        "summary": "Data Engineer with 3 years of experience building reliable batch ETL pipelines, data transformations, and relational database integrations using Python, SQL, Apache Airflow, PostgreSQL, dbt, and Google BigQuery. Strong focus on data quality, query optimization, and automated workflow scheduling.",
        "skills": [
            ("Languages", "Python, SQL, Bash"),
            ("ETL & Orchestration", "Apache Airflow, dbt (data build tool), Python Scripts"),
            ("Data Storage & Warehouses", "Google BigQuery, PostgreSQL, MySQL, AWS S3"),
            ("Tools & Cloud", "Git, Docker, GCP (Cloud Storage, Composer), Linux, JIRA")
        ],
        "experience": [
            {
                "title": "Data Engineer",
                "company": "Genpact, Hyderabad, India",
                "period": "Sep 2021 – Present",
                "bullets": [
                    "Developed and maintained 40+ Airflow DAGs extracting data from REST APIs, SFTP servers, and SQL databases into Google BigQuery.",
                    "Built data transformation models using dbt, enabling self-service analytics for business intelligence teams.",
                    "Optimized SQL queries and table partitioning/clustering in BigQuery, lowering monthly query cost by 30%.",
                    "Wrote automated Python data validation checks ensuring schema adherence before warehouse ingestion."
                ]
            }
        ],
        "projects": [
            {
                "name": "E-Commerce Customer Analytics Data Warehouse",
                "tech": "Python, Airflow, BigQuery, dbt, GCP",
                "bullets": [
                    "Built an ETL pipeline aggregating order sales, user website clicks, and marketing campaign metrics into BigQuery for customer lifetime value (LTV) reporting."
                ]
            },
            {
                "name": "Automated Logistics Tracking Pipeline",
                "tech": "Python, PostgreSQL, Airflow, Docker",
                "bullets": [
                    "Created scheduled data pipelines gathering shipment status logs from 3rd party carrier APIs and updating internal PostgreSQL tables."
                ]
            }
        ],
        "education": [
            "B.E. in Information Technology | Osmania University, Hyderabad, India | Graduated: Jun 2021 | CGPA: 8.1/10.0"
        ],
        "certifications": [
            "Google Professional Data Engineer Certification"
        ],
        "achievements": [
            "Recognized for top data reliability maintenance score across team in 2023."
        ]
    },
    # 24 - Data Engineer
    {
        "num": "24",
        "name": "Deepak Verma",
        "role": "Associate Data Engineer",
        "summary": "Junior Data Engineer with 1 year of experience processing large datasets, writing SQL queries, and developing PySpark scripts. Hands-on experience with Hadoop ecosystem, ETL concepts, relational databases, and AWS S3 storage.",
        "skills": [
            ("Languages", "Python, SQL"),
            ("Big Data & ETL", "PySpark, Apache Spark, Hadoop, HDFS, Hive"),
            ("Databases", "MySQL, PostgreSQL, AWS S3"),
            ("Tools", "Git, Jupyter, Linux Command Line, VS Code")
        ],
        "experience": [
            {
                "title": "Associate Data Engineer",
                "company": "Wipro, Chennai, India",
                "period": "Jul 2023 – Present",
                "bullets": [
                    "Wrote PySpark data processing scripts to clean, transform, and aggregate raw CSV files stored in AWS S3 buckets.",
                    "Created Hive external tables and wrote SQL analytical queries for monthly operations reporting.",
                    "Monitored daily ETL batch jobs, investigating pipeline execution failures and fixing data parsing errors."
                ]
            }
        ],
        "projects": [
            {
                "name": "IoT Telemetry Batch Processing System",
                "tech": "PySpark, AWS S3, MySQL, Python",
                "bullets": [
                    "Built a PySpark batch script reading millions of sensor log records from AWS S3, calculating hourly averages, and writing results to MySQL."
                ]
            },
            {
                "name": "Sales Data Cleaning & Normalization Tool",
                "tech": "Python, Pandas, SQL, SQLite",
                "bullets": [
                    "Developed a utility script normalizing raw multi-currency sales files into standardized relational database schema formats."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | SASTRA University, Thanjavur, India | Graduated: May 2023 | CGPA: 8.2/10.0"
        ],
        "certifications": [
            "Coursera Data Engineering Foundations Specialization"
        ],
        "achievements": [
            "Graduated in Top 10% of Computer Science Class at SASTRA."
        ]
    },

    # 25 - Data Analyst
    {
        "num": "25",
        "name": "Pooja Sundaram",
        "role": "Lead Data Analyst",
        "summary": "Senior Data Analyst with 4+ years of experience transforming raw enterprise data into actionable business insights, interactive dashboards, and executive reports. Expertise in SQL, Python (Pandas/NumPy), Tableau, Power BI, Statistical Modeling, and A/B Testing.",
        "skills": [
            ("BI & Visualization", "Tableau Desktop/Server, Power BI, Looker, Excel (Macros, VBA, Pivot Tables)"),
            ("Data Analysis & Scripting", "SQL (PostgreSQL, Snowflake, MySQL), Python (Pandas, NumPy, Matplotlib, Seaborn), R"),
            ("Statistical Methods", "A/B Testing, Hypothesis Testing, Linear Regression, Cohort Analysis, Segmentation"),
            ("Business Domains", "E-Commerce, Retail, Customer Retention, Financial Metrics (CAC, LTV, Churn)")
        ],
        "experience": [
            {
                "title": "Lead Business Data Analyst",
                "company": "Flipkart, Bangalore, India",
                "period": "Jan 2022 – Present",
                "bullets": [
                    "Built 15+ executive Tableau dashboards tracking daily GMV, user conversion funnels, and seller performance metrics for senior management.",
                    "Designed and analyzed A/B test experiments for checkout page redesigns, driving a 4.2% increase in completed user purchase transactions.",
                    "Executed cohort analysis and customer lifetime value (LTV) segmentation using SQL and Python, identifying key churn indicators.",
                    "Mentored 3 junior analysts in SQL optimization, dashboard formatting, and data storytelling principles."
                ]
            },
            {
                "title": "Data Analyst",
                "company": "LatentView Analytics, Chennai, India",
                "period": "Jul 2019 – Dec 2021",
                "bullets": [
                    "Queried complex relational database tables in SQL Server to generate weekly sales forecasting reports.",
                    "Automated daily reporting workflows using Python scripts and Excel VBA, saving 12 manual team hours per week."
                ]
            }
        ],
        "projects": [
            {
                "name": "E-Commerce Customer Churn Analysis & Prevention",
                "tech": "Python, SQL, Tableau, Logistic Regression",
                "bullets": [
                    "Analyzed 500,000 user activity logs to uncover churn patterns, building an interactive Tableau dashboard flagging at-risk accounts."
                ]
            },
            {
                "name": "Supply Chain Fulfillment Delay Optimization",
                "tech": "SQL, Power BI, Python (Pandas)",
                "bullets": [
                    "Constructed a warehouse logistics dashboard pin-pointing shipping bottlenecks across regional distribution hubs."
                ]
            }
        ],
        "education": [
            "B.Sc in Statistics | Loyola College, Chennai, India | Graduated: May 2019 | Marks: 88%"
        ],
        "certifications": [
            "Tableau Desktop Certified Professional",
            "Microsoft Certified: Power BI Data Analyst Associate (PL-300)"
        ],
        "achievements": [
            "Winner of Flipkart Data Storyteller Award 2023."
        ]
    },
    # 26 - Data Analyst
    {
        "num": "26",
        "name": "Suresh Pillai",
        "role": "Data & Business Analyst",
        "summary": "Data Analyst with 2 years of experience analyzing business metrics, writing complex SQL queries, creating interactive Power BI dashboards, and performing statistical exploratory data analysis (EDA) using Python. Adept at translating business requirements into technical analytical solutions.",
        "skills": [
            ("Data Analytics", "Advanced SQL, Python (Pandas, NumPy, Matplotlib), Advanced Excel (VLOOKUP, Index/Match, Pivot Tables)"),
            ("Visualization", "Power BI, Tableau Public"),
            ("Database Systems", "PostgreSQL, MySQL, MS SQL Server"),
            ("Concepts", "EDA, Data Wrangling, KPI Tracking, Trend Analysis")
        ],
        "experience": [
            {
                "title": "Data Analyst",
                "company": "ZS Associates, Pune, India",
                "period": "Jun 2022 – Present",
                "bullets": [
                    "Developed automated Power BI reports connecting to PostgreSQL databases, tracking monthly healthcare sales performance.",
                    "Wrote complex SQL scripts featuring CTEs, Window Functions, and multi-table Joins to extract business KPIs.",
                    "Performed data cleaning and missing value imputation on raw survey datasets using Python Pandas scripts.",
                    "Presented weekly metric dashboards to international client stakeholders."
                ]
            }
        ],
        "projects": [
            {
                "name": "Retail Store Performance Analytics",
                "tech": "Power BI, SQL, Excel, Python",
                "bullets": [
                    "Created a comprehensive Power BI report visualizing sales trends, profit margins, and regional product demand across 50 retail branches."
                ]
            },
            {
                "name": "Customer Feedback Sentiment Dashboard",
                "tech": "Python, SQL, Power BI",
                "bullets": [
                    "Parsed text feedback data using Python, categorizing user reviews into sentiment buckets and displaying insights in Power BI."
                ]
            }
        ],
        "education": [
            "B.Tech in Mechanical Engineering | College of Engineering Pune (COEP), India | Graduated: May 2022 | CGPA: 7.9/10.0"
        ],
        "certifications": [
            "Google Data Analytics Professional Certificate (Coursera)"
        ],
        "achievements": [
            "Automated ZS team reporting task, cutting weekly analysis overhead by 8 hours."
        ]
    },
    # 27 - Data Analyst
    {
        "num": "27",
        "name": "Ananya Sharma",
        "role": "Junior Data Analyst",
        "summary": "Recent Statistics graduate with strong foundation in SQL, Excel, Python, and data visualization. Knowledgeable in exploratory data analysis, data cleaning, and creating clear graphical summaries for decision makers.",
        "skills": [
            ("Tools", "Microsoft Excel (Pivot Tables, Charts, Formulas), Power BI, Google Sheets"),
            ("Programming & Querying", "SQL (MySQL), Python (Pandas, Seaborn, Matplotlib)"),
            ("Soft Skills", "Problem Solving, Data Storytelling, Communication")
        ],
        "experience": [
            {
                "title": "Data Analyst Intern",
                "company": "KPMG India, Gurgaon, India",
                "period": "Jan 2024 – Jun 2024",
                "bullets": [
                    "Assisted senior analysts in querying SQL databases to aggregate financial transaction metrics for quarterly client audits.",
                    "Cleaned, reformatted, and validated incoming client Excel files containing 50,000+ data rows.",
                    "Built basic Power BI charts demonstrating user demographic distributions."
                ]
            }
        ],
        "projects": [
            {
                "name": "Global Tech Salary Insights Dashboard",
                "tech": "Power BI, SQL, Python, Excel",
                "bullets": [
                    "Cleaned and analyzed open-source tech compensation datasets, creating a Power BI dashboard showing salary breakdown by role, experience, and country."
                ]
            },
            {
                "name": "Movie Rating & Box Office Trend Analysis",
                "tech": "Python, Pandas, Seaborn, Jupyter",
                "bullets": [
                    "Conducted EDA on IMDB movie data, uncovering correlations between budget allocations, genre popularity, and box office revenue."
                ]
            }
        ],
        "education": [
            "B.Sc in Mathematics & Statistics | Delhi University (DU), New Delhi, India | Graduated: Jun 2024 | Marks: 84%"
        ],
        "certifications": [
            "Data Analysis with Python (IBM - Coursera)"
        ],
        "achievements": [
            "Head of Data Science Club at Delhi University (2023-2024)."
        ]
    },

    # 28 - DevOps / Cloud Engineer
    {
        "num": "28",
        "name": "Karthik Subramanian",
        "role": "Senior DevOps & Cloud Architect",
        "summary": "Senior DevOps and Infrastructure Architect with 5.5 years of experience building automation pipelines, cloud infrastructure (AWS/Azure), Kubernetes clusters, and CI/CD systems. Expert in Terraform, Ansible, Docker, Jenkins, Prometheus, and Grafana, with a focus on high availability, zero-downtime deployments, and infrastructure cost optimization.",
        "skills": [
            ("Cloud Platforms", "AWS (EC2, EKS, S3, RDS, VPC, IAM, CloudWatch), Azure"),
            ("Containerization & Orchestration", "Kubernetes, Docker, Helm, ArgoCD, OpenShift"),
            ("Infrastructure as Code (IaC)", "Terraform, Ansible, CloudFormation"),
            ("CI/CD Tools", "Jenkins, GitHub Actions, GitLab CI, Harness"),
            ("Monitoring & Logging", "Prometheus, Grafana, ELK Stack (Elasticsearch, Logstash, Kibana), Datadog"),
            ("Scripting & OS", "Bash, Python, Linux (RHEL, Ubuntu), Git")
        ],
        "experience": [
            {
                "title": "Senior Cloud Infrastructure Engineer",
                "company": "Thoughtworks, Bangalore, India",
                "period": "Oct 2021 – Present",
                "bullets": [
                    "Architected multi-region AWS EKS Kubernetes infrastructure serving 50+ microservices with automated cluster autoscaling (Karpenter).",
                    "Automated end-to-end cloud provisioning using Terraform and Ansible, achieving 100% Infrastructure-as-Code coverage across staging and production.",
                    "Designed GitOps continuous delivery pipelines using ArgoCD and GitHub Actions, lowering release failure rates by 70%.",
                    "Optimized AWS cloud infrastructure usage by implementing spot instance clusters and auto-parking schedules, cutting cloud costs by $120,000 annually."
                ]
            },
            {
                "title": "DevOps Engineer",
                "company": "Capgemini, Pune, India",
                "period": "Jul 2018 – Sep 2021",
                "bullets": [
                    "Managed Jenkins CI/CD automation pipelines for building and deploying Java/Node.js web applications.",
                    "Configured centralized system monitoring and logging using Prometheus and Grafana dashboards."
                ]
            }
        ],
        "projects": [
            {
                "name": "Zero-Downtime E-Commerce Deployment Architecture",
                "tech": "AWS EKS, Terraform, ArgoCD, Helm, Prometheus",
                "bullets": [
                    "Built a progressive blue/green deployment setup on Kubernetes managed via ArgoCD Rollouts with automated rollback triggers on metric spikes."
                ]
            },
            {
                "name": "Automated Security & Compliance Scanning Pipeline",
                "tech": "Trivy, SonarQube, Terraform, GitHub Actions",
                "bullets": [
                    "Integrated container vulnerability scanning and IaC security checks into CI/CD pipelines, blocking insecure commits before merge."
                ]
            }
        ],
        "education": [
            "B.Tech in Computer Science | PSG College of Technology, Coimbatore, India | Graduated: May 2018 | CGPA: 8.9/10.0"
        ],
        "certifications": [
            "Certified Kubernetes Administrator (CKA)",
            "AWS Certified DevOps Engineer – Professional",
            "HashiCorp Certified: Terraform Associate"
        ],
        "achievements": [
            "Speaker at DevOps India Summit 2023 on 'Cost-Effective Kubernetes at Scale'."
        ]
    },
    # 29 - DevOps / Cloud Engineer
    {
        "num": "29",
        "name": "Manish Reddy",
        "role": "Cloud DevOps Engineer",
        "summary": "DevOps Engineer with 3 years of hands-on experience maintaining cloud infrastructure on Microsoft Azure, containerizing applications with Docker, writing Terraform IaC modules, and automating deployment workflows with GitHub Actions and Azure DevOps.",
        "skills": [
            ("Cloud", "Microsoft Azure (AKS, Virtual Machines, Azure Blob, VNet, App Services)"),
            ("Containers & Orchestration", "Docker, Docker Compose, Kubernetes (AKS), Helm"),
            ("CI/CD & IaC", "GitHub Actions, Azure DevOps Pipelines, Terraform"),
            ("Monitoring & Tools", "Azure Monitor, Log Analytics, Git, Bash, Linux")
        ],
        "experience": [
            {
                "title": "Cloud Operations Engineer",
                "company": "NTT DATA, Hyderabad, India",
                "period": "Aug 2021 – Present",
                "bullets": [
                    "Maintained Azure Kubernetes Service (AKS) clusters, configuring horizontal pod autoscalers (HPA) and ingress controllers.",
                    "Wrote reusable Terraform scripts to automate resource group, virtual network, and database provisioning on Azure.",
                    "Built CI/CD automation workflows in GitHub Actions for building Docker images and pushing them to Azure Container Registry (ACR).",
                    "Monitored server metrics using Azure Monitor and set up alerts for high CPU/Memory threshold breaches."
                ]
            }
        ],
        "projects": [
            {
                "name": "Automated Cloud Sandbox Provisioner",
                "tech": "Terraform, Azure DevOps, Azure Cloud, Bash",
                "bullets": [
                    "Created an automated self-service pipeline enabling developers to trigger temporary isolated Azure testing environments."
                ]
            },
            {
                "name": "Microservices Container Migration",
                "tech": "Docker, AKS, Helm, GitHub Actions",
                "bullets": [
                    "Containerized 6 legacy backend services into lightweight Docker containers and deployed them onto AKS using Helm charts."
                ]
            }
        ],
        "education": [
            "B.Tech in Electronics & Communication | NIT Warangal, India | Graduated: May 2021 | CGPA: 8.3/10.0"
        ],
        "certifications": [
            "Microsoft Certified: Azure Administrator Associate (AZ-104)",
            "Certified Kubernetes Application Developer (CKAD)"
        ],
        "achievements": [
            "Achieved 99.9% uptime metric across managed Azure AKS environments."
        ]
    },
    # 30 - DevOps / Cloud Engineer
    {
        "num": "30",
        "name": "Ganesh Kulkarni",
        "role": "Junior DevOps Engineer",
        "summary": "Junior DevOps Engineer with 1 year of experience administering Linux servers, configuring Docker containers, writing Bash scripts, and setting up basic CI/CD build jobs in Jenkins. Dedicated to automation, continuous integration, and infrastructure reliability.",
        "skills": [
            ("Operating Systems", "Linux (Ubuntu, CentOS), Shell Scripting (Bash)"),
            ("DevOps Tools", "Docker, Jenkins, Git, GitHub"),
            ("Cloud", "AWS (EC2, S3, IAM, Security Groups)"),
            ("Networking & Basics", "TCP/IP, SSH, Nginx, Systemd")
        ],
        "experience": [
            {
                "title": "Junior DevOps Engineer",
                "company": "Cybage Software, Pune, India",
                "period": "Jul 2023 – Present",
                "bullets": [
                    "Configured Jenkins pipeline jobs to automate build and execution of unit tests for Java web applications.",
                    "Created Dockerfiles and Docker Compose files to standardize developer local environment setups.",
                    "Managed AWS EC2 instances, security groups, and IAM permission policies under supervision.",
                    "Wrote Bash scripts to automate daily server log rotation and backup tasks."
                ]
            }
        ],
        "projects": [
            {
                "name": "Automated Web Server Setup Script",
                "tech": "Bash, Nginx, AWS EC2, Git",
                "bullets": [
                    "Developed a shell script that automatically provisions a fresh Ubuntu EC2 instance, installs Nginx, configures SSL, and deploys static web content."
                ]
            },
            {
                "name": "Dockerized Multi-Container Application Pipeline",
                "tech": "Docker, Jenkins, GitHub, Python Flask, PostgreSQL",
                "bullets": [
                    "Configured a Jenkins CI build that pulls Flask code, builds Docker container images, and executes containerized integration tests."
                ]
            }
        ],
        "education": [
            "B.E. in Computer Science | Savitribai Phule Pune University, India | Graduated: Jun 2023 | CGPA: 7.8/10.0"
        ],
        "certifications": [
            "AWS Certified Cloud Practitioner"
        ],
        "achievements": [
            "Built automated setup script adopted by Cybage QA team."
        ]
    }
]

def add_heading_with_bottom_border(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(31, 78, 121) # Deep navy blue color
    
    # Add bottom border XML
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8') # 1 pt border
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_docx_resume(data, email, phone):
    doc = Document()
    
    # Page Margins: 0.6 inch top/bottom, 0.75 inch left/right
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Top Identifier (Resume XX)
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = Pt(0)
    p_num.paragraph_format.space_after = Pt(2)
    run_num = p_num.add_run(f"Resume {data['num']}")
    run_num.font.size = Pt(9)
    run_num.font.name = "Calibri"
    run_num.font.color.rgb = RGBColor(128, 128, 128)
    
    # Name Header
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(data['name'].upper())
    run_name.bold = True
    run_name.font.size = Pt(20)
    run_name.font.name = "Calibri"
    run_name.font.color.rgb = RGBColor(31, 78, 121)
    
    # Role Title
    p_role = doc.add_paragraph()
    p_role.paragraph_format.space_before = Pt(0)
    p_role.paragraph_format.space_after = Pt(4)
    p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_role = p_role.add_run(data['role'])
    run_role.bold = True
    run_role.font.size = Pt(12)
    run_role.font.name = "Calibri"
    run_role.font.color.rgb = RGBColor(89, 89, 89)
    
    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(8)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    linkedin = f"linkedin.com/in/{data['name'].lower().replace(' ', '-')}"
    github = f"github.com/{data['name'].lower().replace(' ', '')}-dev"
    contact_str = f"Email: {email}  |  Phone: {phone}  |  LinkedIn: {linkedin}  |  GitHub: {github}"
    run_contact = p_contact.add_run(contact_str)
    run_contact.font.size = Pt(9.5)
    run_contact.font.name = "Calibri"
    
    # Professional Summary
    add_heading_with_bottom_border(doc, "Professional Summary")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(4)
    p_sum.paragraph_format.space_after = Pt(6)
    p_sum.paragraph_format.line_spacing = 1.15
    run_sum = p_sum.add_run(data['summary'])
    run_sum.font.size = Pt(10.5)
    run_sum.font.name = "Calibri"
    
    # Technical Skills
    add_heading_with_bottom_border(doc, "Technical Skills")
    for category, skill_list in data['skills']:
        p_sk = doc.add_paragraph(style='List Bullet')
        p_sk.paragraph_format.space_before = Pt(1)
        p_sk.paragraph_format.space_after = Pt(1)
        
        run_cat = p_sk.add_run(f"{category}: ")
        run_cat.bold = True
        run_cat.font.size = Pt(10)
        run_cat.font.name = "Calibri"
        
        run_val = p_sk.add_run(skill_list)
        run_val.font.size = Pt(10)
        run_val.font.name = "Calibri"
        
    # Work Experience
    add_heading_with_bottom_border(doc, "Work Experience")
    for exp in data['experience']:
        p_exp_head = doc.add_paragraph()
        p_exp_head.paragraph_format.space_before = Pt(6)
        p_exp_head.paragraph_format.space_after = Pt(2)
        
        r_title = p_exp_head.add_run(exp['title'])
        r_title.bold = True
        r_title.font.size = Pt(10.5)
        r_title.font.name = "Calibri"
        
        r_comp = p_exp_head.add_run(f" | {exp['company']}")
        r_comp.italic = True
        r_comp.font.size = Pt(10)
        r_comp.font.name = "Calibri"
        
        r_period = p_exp_head.add_run(f" ({exp['period']})")
        r_period.font.size = Pt(9.5)
        r_period.font.name = "Calibri"
        r_period.font.color.rgb = RGBColor(89, 89, 89)
        
        for b in exp['bullets']:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(1)
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.1
            r_b = p_b.add_run(b)
            r_b.font.size = Pt(10)
            r_b.font.name = "Calibri"
            
    # Projects
    add_heading_with_bottom_border(doc, "Key Projects")
    for proj in data['projects']:
        p_proj_head = doc.add_paragraph()
        p_proj_head.paragraph_format.space_before = Pt(6)
        p_proj_head.paragraph_format.space_after = Pt(2)
        
        r_pname = p_proj_head.add_run(proj['name'])
        r_pname.bold = True
        r_pname.font.size = Pt(10.5)
        r_pname.font.name = "Calibri"
        
        r_ptech = p_proj_head.add_run(f" [{proj['tech']}]")
        r_ptech.italic = True
        r_ptech.font.size = Pt(9.5)
        r_ptech.font.name = "Calibri"
        r_ptech.font.color.rgb = RGBColor(89, 89, 89)
        
        for b in proj['bullets']:
            p_b = doc.add_paragraph(style='List Bullet')
            p_b.paragraph_format.space_before = Pt(1)
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.1
            r_b = p_b.add_run(b)
            r_b.font.size = Pt(10)
            r_b.font.name = "Calibri"
            
    # Education
    add_heading_with_bottom_border(doc, "Education")
    for ed in data['education']:
        p_ed = doc.add_paragraph(style='List Bullet')
        p_ed.paragraph_format.space_before = Pt(2)
        p_ed.paragraph_format.space_after = Pt(2)
        r_ed = p_ed.add_run(ed)
        r_ed.font.size = Pt(10)
        r_ed.font.name = "Calibri"
        
    # Certifications
    if data.get('certifications'):
        add_heading_with_bottom_border(doc, "Certifications")
        for cert in data['certifications']:
            p_c = doc.add_paragraph(style='List Bullet')
            p_c.paragraph_format.space_before = Pt(1)
            p_c.paragraph_format.space_after = Pt(1)
            r_c = p_c.add_run(cert)
            r_c.font.size = Pt(10)
            r_c.font.name = "Calibri"
            
    # Achievements
    if data.get('achievements'):
        add_heading_with_bottom_border(doc, "Achievements")
        for ach in data['achievements']:
            p_a = doc.add_paragraph(style='List Bullet')
            p_a.paragraph_format.space_before = Pt(1)
            p_a.paragraph_format.space_after = Pt(1)
            r_a = p_a.add_run(ach)
            r_a.font.size = Pt(10)
            r_a.font.name = "Calibri"
            
    return doc

# Remove old .txt and .md files to ensure ONLY Word Documents (.docx) remain
old_files = glob.glob(os.path.join(output_dir, "*.txt")) + glob.glob(os.path.join(output_dir, "*.md")) + glob.glob(os.path.join(output_dir, "*.tex"))
for f in old_files:
    try:
        os.remove(f)
    except Exception as e:
        pass

for idx, res in enumerate(resumes_data):
    email = emails[idx % len(emails)]
    phone = phones[idx % len(phones)]
    
    doc = create_docx_resume(res, email, phone)
    filename_docx = f"Resume_{res['num']}_{res['name'].replace(' ', '_')}.docx"
    filepath_docx = os.path.join(output_dir, filename_docx)
    doc.save(filepath_docx)

print("Successfully generated 30 Microsoft Word documents (.docx) in", output_dir)
